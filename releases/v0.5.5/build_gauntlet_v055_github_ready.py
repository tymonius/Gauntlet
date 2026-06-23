from __future__ import annotations

import importlib.util
import json
import math
import shutil
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Tuple
from collections import Counter

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, KeepInFrame

ROOT = Path('/mnt/data/gauntlet_v0_5_5')
SRC_PATH = ROOT / 'build_gauntlet_v055.py'
VERSION = 'v0.5.5'
DATE = 'June 21, 2026'
OWNER = 'Tymon Scott'
COPYRIGHT_SHORT = 'Gauntlet v0.5.5 © 2026 Tymon Scott. All rights reserved. Playtest copy.'
COPYRIGHT_CARD = '© 2026 T. Scott'

spec = importlib.util.spec_from_file_location('gauntlet_src', SRC_PATH)
src = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(src)

PAGE_W, PAGE_H = letter
CARD_W, CARD_H = src.CARD_W, src.CARD_H
LEFT, BOTTOM = src.LEFT, src.BOTTOM


def add_page_number(paragraph) -> None:
    run = paragraph.add_run('Page ')
    run.font.size = Pt(8)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r = OxmlElement('w:r')
    r.append(fldChar1); r.append(instrText); r.append(fldChar2)
    paragraph._p.append(r)


def clear_paragraph(p):
    for child in list(p._p):
        p._p.remove(child)


def patch_docx(path: Path) -> Path:
    doc = Document(path)

    # Footer notice on all sections, with page number retained.
    for sec in doc.sections:
        footer = sec.footer
        p = footer.paragraphs[0]
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(COPYRIGHT_SHORT)
        r.font.size = Pt(7.2)
        if len(footer.paragraphs) < 2:
            p2 = footer.add_paragraph()
        else:
            p2 = footer.paragraphs[1]
            clear_paragraph(p2)
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number(p2)

    # Cover-page visible notice.
    # Add it after the first rule box rather than trying to anchor to the bottom.
    insert_idx = None
    for i, p in enumerate(doc.paragraphs[:18]):
        if 'This is a pre-release ruleset.' in p.text:
            insert_idx = i + 1
            break
    # python-docx has no high-level insert-after, so append a short notice if precise insertion is unavailable.
    if insert_idx is None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(COPYRIGHT_SHORT)
        r.font.size = Pt(8)

    # Append copyright/use appendix.
    doc.add_page_break()
    doc.add_heading('Appendix C. Copyright and use', level=1)
    doc.add_paragraph('Copyright © 2026 Tymon Scott. All rights reserved.')
    doc.add_paragraph(
        'Gauntlet, including its rulebook text, card text, card lists, game materials, print-and-play files, graphic layout, documentation, and related design materials, is owned by Tymon Scott unless otherwise noted.'
    )
    doc.add_paragraph(
        'This playtest package is made available for private review, playtesting, and development reference only. No license is granted to copy, distribute, publish, sell, sublicense, create derivative commercial works from, or otherwise exploit these materials without written permission from the copyright owner.'
    )
    doc.add_paragraph(
        'You may download and print the included materials solely for personal playtesting and feedback related to Gauntlet.'
    )
    doc.add_paragraph(
        'Submitting comments, suggestions, issues, or playtest feedback does not transfer ownership of Gauntlet. By submitting feedback, you grant Tymon Scott permission to use, adapt, and incorporate that feedback into Gauntlet without compensation or obligation.'
    )

    doc.core_properties.author = OWNER
    doc.core_properties.last_modified_by = OWNER
    doc.core_properties.comments = 'Gauntlet playtest materials. All rights reserved.'
    doc.save(path)
    return path


def safe_html(s: str) -> str:
    return src.safe_html(s)


def fit_paragraph(c: canvas.Canvas, html: str, x: float, y: float, w: float, h: float, font_start=7.4, font_min=5.3, leading_ratio=1.16, italic=False) -> float:
    return src.fit_paragraph(c, html, x, y, w, h, font_start, font_min, leading_ratio, italic)


def draw_crop_marks(c, x, y, w, h):
    return src.draw_crop_marks(c, x, y, w, h)


def draw_main_card_with_copyright(c: canvas.Canvas, x: float, y: float, card: Dict[str, Any], deck_name: str|None=None, serial: str='') -> None:
    src.draw_main_card(c, x, y, card, deck_name, serial)
    c.saveState()
    c.setFillColor(colors.black)
    c.setFont('Helvetica', 4.2)
    c.drawCentredString(x + CARD_W / 2, y + 0.074 * inch, COPYRIGHT_CARD)
    c.restoreState()


def draw_heartland_card_with_copyright(c: canvas.Canvas, x: float, y: float, label: str='HEARTLAND', deck_name: str|None=None) -> None:
    c.saveState()
    c.setFillColor(colors.white); c.setStrokeColor(colors.black); c.setLineWidth(1.0)
    c.rect(x, y, CARD_W, CARD_H, fill=1, stroke=1)
    c.setFillColor(colors.HexColor('#D5D5D5'))
    c.rect(x, y + CARD_H - 0.75*inch, CARD_W, 0.75*inch, fill=1, stroke=0)
    c.setFillColor(colors.black); c.setFont('Helvetica-Bold',16)
    c.drawCentredString(x + CARD_W/2, y + CARD_H - 0.47*inch, label)
    fit_paragraph(c, '<b>Your force begins here.</b><br/><br/>An opponent wins by entering this Heartland while it is empty, or by entering and defeating your token here. When you defend here, you have Homeland Advantage and +1.', x + 0.25*inch, y + 0.65*inch, CARD_W - 0.5*inch, CARD_H - 1.65*inch, 10, 7.5, 1.25)
    c.setFont('Helvetica-Bold',7)
    c.drawCentredString(x + CARD_W/2, y + 0.35*inch, 'HEARTLAND IS NOT A TERRITORY')
    c.setFillColor(colors.HexColor('#F2F2F2')); c.rect(x, y, CARD_W, 0.22*inch, fill=1, stroke=0)
    c.setFillColor(colors.black); c.setFont('Helvetica',4.6)
    c.drawString(x + 0.08*inch, y + 0.075*inch, deck_name.upper() if deck_name else 'PLAYTEST')
    c.drawCentredString(x + CARD_W/2, y + 0.075*inch, COPYRIGHT_CARD)
    c.drawRightString(x + CARD_W - 0.08*inch, y + 0.075*inch, VERSION)
    draw_crop_marks(c, x, y, CARD_W, CARD_H)
    c.restoreState()


def draw_aid_card_with_copyright(c: canvas.Canvas, x: float, y: float, title: str, lines: List[str], deck_name: str|None=None) -> None:
    src.draw_aid_card(c, x, y, title, lines, deck_name)
    c.saveState()
    c.setFillColor(colors.black)
    c.setFont('Helvetica', 4.3)
    c.drawCentredString(x + CARD_W / 2, y + 0.06 * inch, COPYRIGHT_CARD)
    c.restoreState()


def fit_landscape_title(c: canvas.Canvas, text: str, x: float, y: float, w: float, font_start=14.0, font_min=8.0):
    fs = font_start
    while stringWidth(text, 'Helvetica-Bold', fs) > w and fs > font_min:
        fs -= 0.3
    c.setFont('Helvetica-Bold', fs)
    c.drawCentredString(x, y, text)


def draw_territory_card_landscape(c: canvas.Canvas, x: float, y: float, territory: Dict[str, str], deck_name: str|None=None, serial: str='') -> None:
    """Draw a portrait poker-card outline with all Territory content rotated for landscape table use."""
    # Physical card border and crop marks remain portrait for cutting.
    c.saveState()
    c.setFillColor(colors.white)
    c.setStrokeColor(colors.black)
    c.setLineWidth(0.9)
    c.rect(x, y, CARD_W, CARD_H, fill=1, stroke=1)
    draw_crop_marks(c, x, y, CARD_W, CARD_H)
    c.restoreState()

    # Rotate the content so the physical card is read landscape after rotation.
    c.saveState()
    c.translate(x + CARD_W, y)
    c.rotate(90)
    # local landscape canvas: width = CARD_H, height = CARD_W
    lw, lh = CARD_H, CARD_W
    is_arena = territory['name'].startswith('Arena:')
    fill = colors.HexColor('#D9D9D9') if not is_arena else colors.HexColor('#C8C8C8')
    footer_h = 0.22 * inch
    header_h = 0.43 * inch
    # Background and header/footer within rotated content
    c.setFillColor(colors.white)
    c.rect(0, 0, lw, lh, fill=1, stroke=0)
    c.setFillColor(fill)
    c.rect(0, lh - header_h, lw, header_h, fill=1, stroke=0)
    c.setFillColor(colors.HexColor('#F2F2F2'))
    c.rect(0, 0, lw, footer_h, fill=1, stroke=0)
    # Labels and title
    c.setFillColor(colors.black)
    c.setFont('Helvetica-Bold', 5.8)
    c.drawString(0.12 * inch, lh - 0.15 * inch, 'ARENA TERRITORY' if is_arena else 'TERRITORY')
    fit_landscape_title(c, territory['name'], lw / 2, lh - 0.31 * inch, lw - 0.36 * inch, 13.0, 7.2)
    # Main text
    body_x = 0.18 * inch
    body_y = footer_h + 0.12 * inch
    body_w = lw - 0.36 * inch
    body_h = lh - header_h - footer_h - 0.22 * inch
    fit_paragraph(c, safe_html(territory['text']), body_x, body_y, body_w, body_h, 8.2, 5.4, 1.12)
    # Footer metadata
    c.setFillColor(colors.black)
    c.setFont('Helvetica', 4.6)
    c.drawString(0.10 * inch, 0.075 * inch, (deck_name.upper() if deck_name else 'MASTER POOL'))
    c.drawCentredString(lw / 2, 0.075 * inch, COPYRIGHT_CARD)
    c.drawRightString(lw - 0.10 * inch, 0.075 * inch, f'{VERSION}{("  "+serial) if serial else ""}')
    c.restoreState()


def page_positions() -> List[Tuple[float, float]]:
    return src.page_positions()


def draw_card_pages(path: Path, items: List[Dict[str, Any]], title_meta: str) -> None:
    c = canvas.Canvas(str(path), pagesize=letter)
    positions = page_positions()
    sheets = math.ceil(len(items)/9)
    for pidx in range(sheets):
        page_items = items[pidx*9:(pidx+1)*9]
        for idx, item in enumerate(page_items):
            x, y = positions[idx]
            kind = item['kind']
            if kind == 'main':
                draw_main_card_with_copyright(c, x, y, item['card'], item.get('deck'), item.get('serial', ''))
            elif kind == 'territory':
                draw_territory_card_landscape(c, x, y, item['territory'], item.get('deck'), item.get('serial', ''))
            elif kind == 'heartland':
                draw_heartland_card_with_copyright(c, x, y, item.get('label', 'HEARTLAND'), item.get('deck'))
            elif kind == 'aid':
                draw_aid_card_with_copyright(c, x, y, item['title'], item['lines'], item.get('deck'))
        c.setFont('Helvetica', 5)
        c.setFillColor(colors.HexColor('#666666'))
        c.drawCentredString(PAGE_W/2, 0.07*inch, f'{title_meta} - sheet {pidx+1} of {sheets} - {COPYRIGHT_SHORT}')
        c.showPage()
    c.save()


def build_four_deck_pdf() -> Path:
    items = []
    for deck_name, deck in src.DECKS.items():
        seq = Counter()
        for name in deck['cards']:
            seq[name] += 1
            items.append({'kind': 'main', 'card': src.CARD_BY_NAME[name], 'deck': deck_name, 'serial': f'{seq[name]}'})
        for t in deck['territories']:
            items.append({'kind': 'territory', 'territory': src.TERRITORY_BY_NAME[t], 'deck': deck_name})
        items.append({'kind': 'heartland', 'deck': deck_name})
        items.append({'kind': 'aid', 'deck': deck_name, 'title': 'TURN SEQUENCE', 'lines': [
            'Capture check.', 'Draw one card.', 'Optional Action before movement.', 'Advance, hold, or withdraw.', 'Optional Action after movement if unused.', 'Cleanup to three cards.'
        ]})
        items.append({'kind': 'aid', 'deck': deck_name, 'title': 'BATTLE & DESTINATIONS', 'lines': [
            'Optional one card from hand.', 'Draw three; select up to one.', 'Reveal; cancel; resolve effects.', 'Roll and determine winner.', 'Loser retreats; attacker may occupy.', 'Hand to GY; battle-draw to discard. Heartland: ties/+1.'
        ]})
    assert len(items) == 144
    path = ROOT / 'Gauntlet_v0.5.5_Four_Deck_Playtest_Set.pdf'
    draw_card_pages(path, items, 'Gauntlet v0.5.5 Four-Deck Playtest Set')
    return path


def build_master_pool_pdf() -> Path:
    items = [{'kind': 'main', 'card': card} for card in src.CARDS]
    items += [{'kind': 'territory', 'territory': t} for t in src.TERRITORIES]
    items.append({'kind': 'heartland', 'label': 'HEARTLAND A'})
    items.append({'kind': 'heartland', 'label': 'HEARTLAND B'})
    assert len(items) == 81
    path = ROOT / 'Gauntlet_v0.5.5_Master_Card_Pool.pdf'
    draw_card_pages(path, items, 'Gauntlet v0.5.5 Master Card Pool')
    return path


def build_playtest_forms_pdf() -> Path:
    # Use source form builder, then stamp copyright with a simple footer overlay by regenerating from existing code pattern would be overkill.
    path = src.build_playtest_forms_pdf()
    # Rebuild in place with copyright footer to avoid external PDF overlay dependencies.
    c = canvas.Canvas(str(path), pagesize=letter)
    for page in range(4):
        margin=0.55*inch
        c.setFont('Helvetica-Bold',18); c.drawString(margin,PAGE_H-margin,'GAUNTLET PLAYTEST RECORD')
        c.setFont('Helvetica-Bold',10); c.drawRightString(PAGE_W-margin,PAGE_H-margin+2,VERSION)
        y=PAGE_H-margin-0.38*inch
        fields=[('Date / location',2.0),('Players',2.0),('Deck A / player',2.4),('Deck B / player',2.4),('First player',1.8),('Winner',1.8),('Ending',2.8),('Elapsed time',1.8),('Player-turns',1.8),('Battles',1.4),('Captures',1.4),('Reshuffles A / B',2.0),('Final Graveyard A / B',2.3)]
        colx=[margin,PAGE_W/2+0.12*inch]
        idx=0
        for row in range(7):
            for col in range(2):
                if idx>=len(fields): break
                label,_=fields[idx]; x=colx[col]
                c.setFont('Helvetica-Bold',7.5); c.drawString(x,y,label.upper())
                c.line(x,y-0.08*inch,x+(PAGE_W/2-margin-0.2*inch),y-0.08*inch)
                idx+=1
            y-=0.42*inch
        y-=0.05*inch
        c.setFont('Helvetica-Bold',9); c.drawString(margin,y,'PLAYER RATINGS (1 = poor, 5 = excellent)')
        y-=0.25*inch
        c.setFont('Helvetica',8)
        for label in ['Enjoyment','Agency','Pacing','Desire to replay']:
            c.drawString(margin,y,label+' - Player A:  1  2  3  4  5        Player B:  1  2  3  4  5')
            y-=0.25*inch
        prompts=['Rules questions or ambiguous interactions','Most valuable card or Territory','Least valuable card or Territory','Did either player feel unable to act? Why?','Was the winner apparent long before the game ended?','Other observations']
        for prompt in prompts:
            c.setFont('Helvetica-Bold',8); c.drawString(margin,y,prompt.upper()); y-=0.13*inch
            lines=2 if prompt!='Other observations' else 4
            for _ in range(lines):
                c.line(margin,y,PAGE_W-margin,y); y-=0.22*inch
            y-=0.07*inch
        c.setFont('Helvetica',5.5)
        c.drawCentredString(PAGE_W/2,0.28*inch,f'Gauntlet {VERSION} - Record sheet {page+1} of 4')
        c.drawCentredString(PAGE_W/2,0.18*inch,COPYRIGHT_SHORT)
        c.showPage()
    c.save()
    return path


def update_json(path: Path) -> Path:
    data = json.loads(path.read_text(encoding='utf-8'))
    data['copyright'] = {
        'notice': 'Copyright © 2026 Tymon Scott. All rights reserved.',
        'owner': 'Tymon Scott',
        'use': 'Provided for private review, playtesting, and development reference only. No license is granted except personal printing for Gauntlet playtesting and feedback.',
        'feedback': 'Submitting feedback grants Tymon Scott permission to use, adapt, and incorporate that feedback into Gauntlet without compensation or obligation.'
    }
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding='utf-8')
    return path


def update_deck_lists(path: Path) -> Path:
    text = path.read_text(encoding='utf-8')
    notice = f'**Copyright:** {COPYRIGHT_SHORT}\n\n'
    if '**Copyright:**' not in text:
        text = text.replace(f'**Date:** {DATE}\n\n', f'**Date:** {DATE}\n\n{notice}')
    path.write_text(text, encoding='utf-8')
    return path


def write_copyright() -> Path:
    text = '''# Copyright and permitted use

Copyright © 2026 Tymon Scott. All rights reserved.

Gauntlet, including its rulebook text, card text, card lists, game materials, print-and-play files, graphic layout, documentation, and related design materials, is owned by Tymon Scott unless otherwise noted.

This repository is made available for private review, playtesting, and development reference only. No license is granted to copy, distribute, publish, sell, sublicense, create derivative commercial works from, or otherwise exploit these materials without written permission from the copyright owner.

You may download and print the included materials solely for personal playtesting and feedback related to Gauntlet.

Submitting comments, suggestions, issues, or playtest feedback does not transfer ownership of Gauntlet. By submitting feedback, you grant Tymon Scott permission to use, adapt, and incorporate that feedback into Gauntlet without compensation or obligation.

All rights reserved.
'''
    path = ROOT / 'COPYRIGHT.md'
    path.write_text(text, encoding='utf-8')
    return path


def write_contributing() -> Path:
    text = '''# Contributing and playtest feedback

Gauntlet is an unpublished playtest project owned by Tymon Scott.

Feedback, rules questions, playtest reports, bug reports, balance suggestions, and wording suggestions are welcome. By submitting comments, issues, pull requests, suggested card text, playtest notes, or other feedback, you grant Tymon Scott permission to use, adapt, modify, publish, and incorporate that feedback into Gauntlet without compensation, credit obligation, or transfer of ownership.

Do not submit material you do not have the right to share. Do not submit third-party copyrighted artwork, card text, game text, lore, or other protected material unless you have permission to do so.

Submitting feedback does not give you ownership of Gauntlet or any portion of the project. No license is granted to copy, redistribute, sell, publish, sublicense, or create derivative commercial works from Gauntlet materials.

For ordinary playtest reports, please include:

- Version number used
- Decks used
- Winner and ending type
- Approximate elapsed time
- Approximate player-turn count
- What felt fun
- What felt slow, confusing, or unfair
- Any card or Territory interaction that required a ruling
'''
    path = ROOT / 'CONTRIBUTING.md'
    path.write_text(text, encoding='utf-8')
    return path


def write_readme() -> Path:
    text = f'''# Gauntlet {VERSION} Print-and-Play Package

Generated {DATE}. GitHub-ready document revision with copyright notices and landscape-oriented Territory card text.

## Files

- `Gauntlet_v0.5.5_Complete_Playtest_Guide.docx` - editable canonical guide.
- `Gauntlet_v0.5.5_Complete_Playtest_Guide.pdf` - print-ready guide.
- `Gauntlet_v0.5.5_Four_Deck_Playtest_Set.pdf` - four complete 30-card decks, twelve Territories, four Heartlands, and eight player aids.
- `Gauntlet_v0.5.5_Master_Card_Pool.pdf` - one copy of all 54 main-deck cards, all 25 Territories, and two Heartlands.
- `Gauntlet_v0.5.5_Playtest_Forms.pdf` - four game record sheets.
- `Gauntlet_v0.5.5_Canonical_Data.json` - editable structured source data.
- `Gauntlet_v0.5.5_Playtest_Deck_Lists.md` - deck manifests.
- `COPYRIGHT.md` - copyright and permitted-use notice.
- `CONTRIBUTING.md` - feedback/contribution terms.

## Printing

Card sheets use exact 2.5 x 3.5 inch poker-card fronts in a 3 x 3 US Letter layout. Print at **Actual Size / 100%**, with scaling disabled. For a durable prototype, cut the fronts and sleeve them in opaque standard poker sleeves with ordinary playing cards or bulk trading cards behind them. This avoids duplex alignment problems.

Territory cards are still cut to standard poker-card size, but their text is printed in landscape orientation so the card can face the controlling player on the board.

The Four-Deck Playtest Set is organized deck-by-deck and labels every card at the bottom for sorting. The Shadow Directorate is the v0.5.5 Attrition trial list; Attrition replaces Court Martial at equal cost.

## Rules changes incorporated

- Heartland terminology.
- Hand-origin battle cards go to the Graveyard.
- Battle-drawn battle cards go to the discard pile.
- Replayed Battle cards normally go to the Graveyard.
- Incomplete draws draw as many cards as possible.
- Attrition uses the locked three-card Battle effect.
- Last Stand: Heartland defenders have Homeland Advantage and gain +1.
- Human playtest errata for Conscription, Watchtower, Fortifications, and Embargo.

## Copyright and use

Gauntlet is currently an unpublished playtest project.

Copyright © 2026 Tymon Scott. All rights reserved.

The materials in this repository are provided for private review and playtesting only. They may not be copied, redistributed, sold, republished, or used to create commercial derivative works without written permission.
'''
    path = ROOT / 'README.md'
    path.write_text(text, encoding='utf-8')
    return path


def write_release_notes() -> Path:
    text = f'''# Gauntlet {VERSION} - Human Playtest Errata

Release date: {DATE}

Copyright © 2026 Tymon Scott. All rights reserved. Playtest copy.

## Rules/card errata

- Reworked Conscription's Battle effect so it functions within battle-draw timing.
- Clarified that Watchtower only reveals attacker hand commitments made during the normal battle commitment step.
- Reworked Fortifications' Battle effect into an immediate defensive +1 and optional extra retreat after a loss.
- Rewrote Embargo so canceled cards return to the appropriate source/destination rather than always entering hand.

## Production/document revision before GitHub push

- Added copyright notices to the guide, PDFs, README, release notes, deck lists, card-sheet metadata, and canonical JSON.
- Added `COPYRIGHT.md` reserving all rights and permitting personal printing only for playtesting and feedback.
- Added `CONTRIBUTING.md` clarifying feedback and playtest-submission terms.
- Updated Territory card fronts so Territory text is printed in landscape orientation for board use.

## Pacing note

No pacing-system rule was added in this patch. Foothold, failed-counterattack capture, and other anti-stalemate options remain under consideration for a later patch.
'''
    path = ROOT / 'Gauntlet_v0.5.5_Release_Notes.md'
    path.write_text(text, encoding='utf-8')
    return path


def make_zip(paths: List[Path]) -> Path:
    zpath = ROOT / 'Gauntlet_v0.5.5_Print_and_Play_Package.zip'
    with zipfile.ZipFile(zpath, 'w', zipfile.ZIP_DEFLATED) as z:
        for p in paths:
            if p.exists():
                z.write(p, arcname=p.name)
    return zpath


def main():
    src.validate()
    outputs: List[Path] = []
    outputs.append(src.export_json())
    outputs[-1] = update_json(outputs[-1])
    outputs.append(src.export_deck_lists_md())
    outputs[-1] = update_deck_lists(outputs[-1])
    outputs.append(src.build_docx())
    outputs[-1] = patch_docx(outputs[-1])
    outputs.append(build_four_deck_pdf())
    outputs.append(build_master_pool_pdf())
    outputs.append(build_playtest_forms_pdf())
    outputs.append(write_readme())
    outputs.append(write_copyright())
    outputs.append(write_contributing())
    outputs.append(write_release_notes())
    # Copy this build script into package as the source generator.
    outputs.append(Path(__file__))
    # PDF version of guide is generated externally by render_docx.py; include after build if present.
    print('\n'.join(str(p) for p in outputs))

if __name__ == '__main__':
    main()
