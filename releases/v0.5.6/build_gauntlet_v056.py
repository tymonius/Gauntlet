from __future__ import annotations

import importlib.util
import json
import math
import shutil
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Tuple
from collections import Counter

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, KeepInFrame

ROOT = Path('/mnt/data/gauntlet_v0_5_6')
SRC_PATH = ROOT / 'build_gauntlet_v055.py'
VERSION = 'v0.5.6'
DATE = 'June 21, 2026'
OWNER = 'Tymon Scott'
COPYRIGHT_SHORT = 'Gauntlet v0.5.6 © 2026 Tymon Scott. All rights reserved. Playtest copy.'
COPYRIGHT_CARD = '© 2026 T. Scott'

spec = importlib.util.spec_from_file_location('gauntlet_src', SRC_PATH)
src = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(src)


def apply_v056_changes() -> None:
    """Mutate the v0.5.5 source data into the v0.5.6 Asset Bank Patch."""
    src.VERSION = VERSION
    src.DATE = DATE
    src.ROOT = ROOT

    by = {c['name']: c for c in src.CARDS}
    def card(name: str, **fields: str) -> None:
        by[name].update(fields)

    # Asset banking terminology.
    card('Arcane Knowledge', action='Bank Arcane Knowledge as an Asset. Once per turn during a battle, Arcane Knowledge may use the eligible Battle effect of one other card you played in that battle.')
    card('Attrition', action="Bank Attrition as an Asset. Whenever your opponent loses a battle against you, each card they played from their battle draw goes to their Graveyard instead of their discard pile.")
    card('Brothers in Arms', action='Bank Brothers in Arms as an Asset. Whenever you play both a card from your hand and a card from your battle draw in the same battle, gain advantage. If you win, place the hand-origin card in your discard pile instead of your Graveyard.')
    card('Conscription', action='Draw one card. You may immediately play one card from your hand whose Action banks it as an Asset.')
    card('Counterintelligence', action='Bank Counterintelligence as an Asset. Opposing effects cannot look at, reveal, or require you to reveal your hand, battle draw, face-down Battle cards, or face-down Territories. This does not prevent revelation required by the normal rules of battle or Territory exploration.')
    card('Decoys', action='Bank Decoys as an Asset. If an opposing effect would cause one or more of your other Assets to leave play, you may discard Decoys. If you do, choose one affected Asset; it remains in play.')
    card('Entrenchment', action='Bank Entrenchment as an Asset. When your opponent advances onto a Territory adjacent to your token, their movement ends and they cannot play an Action card after movement that turn.')
    card('Fealty', action='Bank Fealty as an Asset. Opposing card effects cannot give you disadvantage.')
    card('Fortifications', action='Bank Fortifications as an Asset. When defending, you may play up to two cards from your battle draw instead of one.')
    card('Illegal Occupation', action='Bank Illegal Occupation as an Asset. While your opponent occupies a Territory you control without controlling it, their Assets are inactive.')
    card('Liberation', action='Bank Liberation as an Asset. Whenever you win a counterattack against an opponent occupying a Territory you control, draw one card and you may immediately play one additional Action card.')
    card('Militias', action='Bank Militias as an Asset. During battles on a Territory you control, your opponent gains disadvantage.')
    card('Patriotism', action='Bank Patriotism as an Asset. During battles on a Territory you control, the first +1 or advantage granted by one of your Battle cards is doubled.')
    card('Protracted Siege', action='Bank Protracted Siege as an Asset. Whenever your opponent occupies a Territory you control, they do not capture it at the beginning of their next turn. If they still occupy it at the beginning of the following turn, they capture it normally.')
    card('Resistance', action='Bank Resistance as an Asset. When counterattacking an opponent occupying a Territory you control, draw two additional battle cards before choosing which battle-drawn card to play.', battle='If you are counterattacking an opponent occupying a Territory you control, gain advantage. If you win, bank Resistance as an Asset instead of its normal destination.', reminder='If your Asset bank is at its limit, you may discard one Asset to make room; otherwise Resistance goes to its normal destination.')
    card('Rousing Speech', action='Bank Rousing Speech as an Asset. Whenever your opponent banks a new Asset, you may draw one card, then discard one card.', reminder='Turning an existing Asset face up does not trigger Rousing Speech.')
    card('Spies', action='Bank Spies as an Asset. Your opponent keeps their hand face up.')
    card('Strategic Withdrawal', action='Return one Asset you control from your Asset bank to your hand. If you do, gain one additional movement this turn.')
    card('Tyranny', action="Bank Tyranny as an Asset. Once per battle, the first opposing Battle card that would add to your opponent's battle total or grant them advantage has those portions of its effect ignored.")
    card('Valor', action='Bank Valor as an Asset. Whenever you lose a battle, draw one card after retreating.')
    card('War Crimes', action='Bank War Crimes as an Asset. Whenever your opponent loses a battle against you, they cannot benefit from effects they control that trigger because they lost or retreated. Cards they played during that battle must go to their Graveyard unless an opposing effect moves them elsewhere.')

    # Condition terminology and delayed-action clarification.
    card('Armistice', action="Play Armistice as a Condition. For the rest of this turn and throughout your opponent's next turn, neither player may initiate a battle. Discard Armistice at the end of that turn.")
    card('Assimilation', action='Play Assimilation as a Condition. If you win your next battle this turn as the attacker on a Territory your opponent controls, immediately capture that Territory instead of occupying it. If another effect would delay capture, reduce that delay by one round instead. Discard Assimilation at the end of the turn.')
    card('Blockade', action="Play Blockade as a Condition affecting your opponent. At the beginning of each of their turns, after their normal draw, if they have more than one card in hand, they discard one at random. When they win a battle, discard Blockade.", battle="Your opponent may discard one card from their hand. If they do not, each card they played from their battle draw has no effect during this battle. If they lose, play Blockade as a Condition affecting them instead of its normal destination.")
    card('Capital Gains', action='As an additional cost to play Capital Gains, discard one other card from your hand. Play Capital Gains as a Condition. At the beginning of your next turn, draw three cards instead of your normal draw, then discard Capital Gains.')
    card('Capital Punishment', action="Play Capital Punishment as a Condition. Choose one opposing Asset. If you defeat that opponent in battle this turn, place the chosen Asset in its owner's Graveyard. Discard Capital Punishment at the end of the turn.")
    card('Court Martial', action="Play Court Martial as a Condition affecting your opponent. After their next battle, discard Court Martial. If they lost that battle, they must retreat one additional tile.")
    card('Embargo', action="Play Embargo as a Condition affecting your opponent. At the beginning of their next turn, after their normal draw, they discard one card from their hand at random. Then discard Embargo.")
    card('Palisade Wall', action="Play Palisade Wall as a Condition. During the next battle in which you are the defender, your opponent's Assets are inactive. Discard Palisade Wall after that battle.")
    card('Protracted Siege', battle="If you are defending a Territory you control and lose, play Protracted Siege as a Condition affecting your opponent instead of its normal destination. They do not capture the Territory occupied during this battle at the beginning of their next turn. If they still occupy it at the beginning of the following turn, they capture it normally, then discard Protracted Siege. If they cease occupying it first, discard Protracted Siege.")
    card('Redemption', action='Play Redemption as a Condition. The next time an opposing effect causes one or more of your cards to enter your discard pile, choose one of those cards and return it to your hand after that effect resolves. Then discard Redemption.')
    card('Reinforcements', action='Play Reinforcements as a Condition. During your next turn, you may play one additional Action card. Discard Reinforcements at the end of that turn.')
    card('Sabotage', action='Play Sabotage as a Condition attached to one opposing Asset. Turn that Asset face down until the start of your next turn, then discard Sabotage.')
    card('Scorched Earth', action="Play Scorched Earth as a Condition affecting your opponent. After their next battle, if they won, discard Scorched Earth. If they lost, they choose one Asset they control and discard it, and they cannot bank Assets during their next turn. Discard Scorched Earth at the end of that turn.", battle='If your opponent loses this battle, they choose one Asset they control and discard it. Play Scorched Earth as a Condition affecting them instead of its normal destination. They cannot bank Assets during their next turn. Discard Scorched Earth at the end of that turn.')
    card('Shock and Awe', action='Play Shock and Awe as a Condition. If you win your next battle this turn as the attacker on a Territory your opponent controls, immediately capture that Territory instead of occupying it, then you may advance one additional tile. If another effect would delay the capture, reduce that delay by one round instead; you may advance only if the Territory is captured. Discard Shock and Awe at the end of the turn.')
    card('Stand Ground', action='Play Stand Ground as a Condition. Until the start of your next turn, your token cannot be moved by opposing card effects. Then discard Stand Ground.')
    card('Supplies', action='Play Supplies as a Condition. At the beginning of your next turn, draw two additional cards, then discard Supplies.')
    card('Tariffs', action='Play Tariffs as a Condition. Draw two cards. You may immediately play one additional Action card. At the beginning of your next turn, skip your normal draw, then discard Tariffs.', reminder='Multiple Tariffs may chain, but the skipped-draw penalty does not stack.')
    card('Treason', action="Play Treason as a Condition affecting your opponent. During their next battle, after Battle cards are revealed, choose one eligible card they committed from hand. That card has no effect; place it in its owner's discard pile immediately. Treason uses that card's Battle effect for you. Then discard Treason.")

    # Global wording cleanup inside remaining card/territory text.
    for c in src.CARDS:
        for k in ('action','battle','reminder'):
            if k in c:
                c[k] = c[k].replace('Assets area', 'Asset bank').replace('play Assets', 'bank Assets')
    for t in src.TERRITORIES:
        t['text'] = t['text'].replace('Assets area', 'Asset bank')

    src.VERSION_HISTORY.append({
        'version': 'v0.5.6',
        'name': 'Asset Bank Patch',
        'changes': [
            'Renamed the persistent Asset area to the Asset bank and standardized the language of banking an Asset.',
            'Asset bank limit now equals the number of Territories a player controls; occupied Territories do not count until captured.',
            'If a player controls more Assets than their current Asset bank limit, they immediately discard Assets until within the limit.',
            'Clarified Action-card destinations: Actions that are not banked as Assets, discarded after resolving, or placed as Overlays/Territories are played as Conditions and discarded when directed by their text.',
            'Updated affected card wording to identify when cards are banked as Assets or played as Conditions.'
        ]
    })
    src.CARD_BY_NAME = {c['name']: c for c in src.CARDS}
    src.TERRITORY_BY_NAME = {t['name']: t for t in src.TERRITORIES}


PAGE_W, PAGE_H = letter
CARD_W, CARD_H = src.CARD_W, src.CARD_H
LEFT, BOTTOM = src.LEFT, src.BOTTOM


def add_page_number(paragraph) -> None:
    run = paragraph.add_run('Page ')
    run.font.size = Pt(8)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r = OxmlElement('w:r')
    r.append(fldChar1); r.append(instrText); r.append(fldChar2)
    paragraph._p.append(r)


def clear_paragraph(p):
    for child in list(p._p):
        p._p.remove(child)


def patch_docx(path: Path) -> Path:
    doc = Document(path)

    # Footer notice on all sections, with page number retained.
    for sec in doc.sections:
        footer = sec.footer
        p = footer.paragraphs[0]
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(COPYRIGHT_SHORT)
        r.font.size = Pt(7.2)
        if len(footer.paragraphs) < 2:
            p2 = footer.add_paragraph()
        else:
            p2 = footer.paragraphs[1]
            clear_paragraph(p2)
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number(p2)

    # Cover-page visible notice.
    # Add it after the first rule box rather than trying to anchor to the bottom.
    insert_idx = None
    for i, p in enumerate(doc.paragraphs[:18]):
        if 'This is a pre-release ruleset.' in p.text:
            insert_idx = i + 1
            break
    # python-docx has no high-level insert-after, so append a short notice if precise insertion is unavailable.
    if insert_idx is None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(COPYRIGHT_SHORT)
        r.font.size = Pt(8)

    # v0.5.6 terminology and rules patch in generated rulebook.
    replacements = {
        'A discard pile, face-up Graveyard, Assets area with three spaces, and Conditions area for each player.': 'A discard pile, face-up Graveyard, Asset bank, and Conditions area for each player.',
        'Unless it becomes an Asset, Condition, Overlay, Territory, or says otherwise.': 'Unless it is banked as an Asset, played as a Condition, becomes an Overlay or Territory, or says otherwise.',
        'Resolve remaining after-battle effects, including draws, recovery, and cards that become Assets or Conditions.': 'Resolve remaining after-battle effects, including draws, recovery, and cards that are banked as Assets or played as Conditions.',
        'Assets, Conditions, Overlays, the new Territory, Homeland Advantage, and Last Stand are checked again.': 'Asset bank, Conditions, Overlays, the new Territory, Homeland Advantage, and Last Stand are checked again.',
        'An Asset is a persistent card in your Assets area. You may control no more than three.': 'An Asset is a persistent card in your Asset bank. Your Asset bank limit equals the number of Territories you control. Occupied Territories do not count until captured; Heartlands are not Territories.',
        'To play an Asset into a full area, you may discard one Asset you control to make room.': 'To bank an Asset when your Asset bank is at its limit, you may discard one Asset you control to make room.',
        'If an effect would place a card into your full Assets area, you may discard an Asset to make room; otherwise the incoming card goes to its normal destination.': 'If an effect would bank an Asset while your Asset bank is at its limit, you may discard an Asset to make room; otherwise the incoming card goes to its normal destination.',
        'A face-down Asset remains controlled and occupies a slot, but its printed effects are inactive.': 'A face-down Asset remains controlled and occupies space in your Asset bank, but its printed effects are inactive.',
        'Removed Assets normally go to the discard pile unless an effect says otherwise.': 'Removed Assets normally go to the discard pile unless an effect says otherwise. If your Asset bank limit decreases below the number of Assets you control, immediately choose and discard Assets until you are within the limit.',
        'Conditions do not occupy Asset slots.': 'Conditions do not occupy the Asset bank.',
        'A Condition is a face-up card placed on a player until its printed trigger, duration, or removal condition ends.': 'A Condition is a face-up Action card with a delayed or ongoing effect. An Action that is not banked as an Asset, discarded after resolving, or placed as an Overlay or Territory is played as a Condition and remains until its text tells you to discard it.',
    }
    def replace_runs(paragraph):
        for run in paragraph.runs:
            for a, b in replacements.items():
                if a in run.text:
                    run.text = run.text.replace(a, b)
            run.text = run.text.replace('Assets area', 'Asset bank')
            run.text = run.text.replace('full Asset bank', 'Asset bank at its limit')
            run.text = run.text.replace('Asset slots', 'Asset bank')
    for p0 in doc.paragraphs:
        replace_runs(p0)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p0 in cell.paragraphs:
                    replace_runs(p0)

    # Append copyright/use appendix.
    doc.add_page_break()
    doc.add_heading('Appendix C. Copyright and use', level=1)
    doc.add_paragraph('Copyright © 2026 Tymon Scott. All rights reserved.')
    doc.add_paragraph(
        'Gauntlet, including its rulebook text, card text, card lists, game materials, print-and-play files, graphic layout, documentation, and related design materials, is owned by Tymon Scott unless otherwise noted.'
    )
    doc.add_paragraph(
        'This playtest package is made available for private review, playtesting, and development reference only. No license is granted to copy, distribute, publish, sell, sublicense, create derivative commercial works from, or otherwise exploit these materials without written permission from the copyright owner.'
    )
    doc.add_paragraph(
        'You may download and print the included materials solely for personal playtesting and feedback related to Gauntlet.'
    )
    doc.add_paragraph(
        'Submitting comments, suggestions, issues, or playtest feedback does not transfer ownership of Gauntlet. By submitting feedback, you grant Tymon Scott permission to use, adapt, and incorporate that feedback into Gauntlet without compensation or obligation.'
    )

    doc.core_properties.author = OWNER
    doc.core_properties.last_modified_by = OWNER
    doc.core_properties.comments = 'Gauntlet playtest materials. All rights reserved.'
    doc.save(path)
    return path


def safe_html(s: str) -> str:
    return src.safe_html(s)


def fit_paragraph(c: canvas.Canvas, html: str, x: float, y: float, w: float, h: float, font_start=7.4, font_min=5.3, leading_ratio=1.16, italic=False) -> float:
    return src.fit_paragraph(c, html, x, y, w, h, font_start, font_min, leading_ratio, italic)


def draw_crop_marks(c, x, y, w, h):
    return src.draw_crop_marks(c, x, y, w, h)


def draw_main_card_with_copyright(c: canvas.Canvas, x: float, y: float, card: Dict[str, Any], deck_name: str|None=None, serial: str='') -> None:
    src.draw_main_card(c, x, y, card, deck_name, serial)
    c.saveState()
    c.setFillColor(colors.black)
    c.setFont('Helvetica', 4.2)
    c.drawCentredString(x + CARD_W / 2, y + 0.074 * inch, COPYRIGHT_CARD)
    c.restoreState()


def draw_heartland_card_with_copyright(c: canvas.Canvas, x: float, y: float, label: str='HEARTLAND', deck_name: str|None=None) -> None:
    c.saveState()
    c.setFillColor(colors.white); c.setStrokeColor(colors.black); c.setLineWidth(1.0)
    c.rect(x, y, CARD_W, CARD_H, fill=1, stroke=1)
    c.setFillColor(colors.HexColor('#D5D5D5'))
    c.rect(x, y + CARD_H - 0.75*inch, CARD_W, 0.75*inch, fill=1, stroke=0)
    c.setFillColor(colors.black); c.setFont('Helvetica-Bold',16)
    c.drawCentredString(x + CARD_W/2, y + CARD_H - 0.47*inch, label)
    fit_paragraph(c, '<b>Your force begins here.</b><br/><br/>An opponent wins by entering this Heartland while it is empty, or by entering and defeating your token here. When you defend here, you have Homeland Advantage and +1.', x + 0.25*inch, y + 0.65*inch, CARD_W - 0.5*inch, CARD_H - 1.65*inch, 10, 7.5, 1.25)
    c.setFont('Helvetica-Bold',7)
    c.drawCentredString(x + CARD_W/2, y + 0.35*inch, 'HEARTLAND IS NOT A TERRITORY')
    c.setFillColor(colors.HexColor('#F2F2F2')); c.rect(x, y, CARD_W, 0.22*inch, fill=1, stroke=0)
    c.setFillColor(colors.black); c.setFont('Helvetica',4.6)
    c.drawString(x + 0.08*inch, y + 0.075*inch, deck_name.upper() if deck_name else 'PLAYTEST')
    c.drawCentredString(x + CARD_W/2, y + 0.075*inch, COPYRIGHT_CARD)
    c.drawRightString(x + CARD_W - 0.08*inch, y + 0.075*inch, VERSION)
    draw_crop_marks(c, x, y, CARD_W, CARD_H)
    c.restoreState()


def draw_aid_card_with_copyright(c: canvas.Canvas, x: float, y: float, title: str, lines: List[str], deck_name: str|None=None) -> None:
    src.draw_aid_card(c, x, y, title, lines, deck_name)
    c.saveState()
    c.setFillColor(colors.black)
    c.setFont('Helvetica', 4.3)
    c.drawCentredString(x + CARD_W / 2, y + 0.06 * inch, COPYRIGHT_CARD)
    c.restoreState()


def fit_landscape_title(c: canvas.Canvas, text: str, x: float, y: float, w: float, font_start=14.0, font_min=8.0):
    fs = font_start
    while stringWidth(text, 'Helvetica-Bold', fs) > w and fs > font_min:
        fs -= 0.3
    c.setFont('Helvetica-Bold', fs)
    c.drawCentredString(x, y, text)


def draw_territory_card_landscape(c: canvas.Canvas, x: float, y: float, territory: Dict[str, str], deck_name: str|None=None, serial: str='') -> None:
    """Draw a portrait poker-card outline with all Territory content rotated for landscape table use."""
    # Physical card border and crop marks remain portrait for cutting.
    c.saveState()
    c.setFillColor(colors.white)
    c.setStrokeColor(colors.black)
    c.setLineWidth(0.9)
    c.rect(x, y, CARD_W, CARD_H, fill=1, stroke=1)
    draw_crop_marks(c, x, y, CARD_W, CARD_H)
    c.restoreState()

    # Rotate the content so the physical card is read landscape after rotation.
    c.saveState()
    c.translate(x + CARD_W, y)
    c.rotate(90)
    # local landscape canvas: width = CARD_H, height = CARD_W
    lw, lh = CARD_H, CARD_W
    is_arena = territory['name'].startswith('Arena:')
    fill = colors.HexColor('#D9D9D9') if not is_arena else colors.HexColor('#C8C8C8')
    footer_h = 0.22 * inch
    header_h = 0.43 * inch
    # Background and header/footer within rotated content
    c.setFillColor(colors.white)
    c.rect(0, 0, lw, lh, fill=1, stroke=0)
    c.setFillColor(fill)
    c.rect(0, lh - header_h, lw, header_h, fill=1, stroke=0)
    c.setFillColor(colors.HexColor('#F2F2F2'))
    c.rect(0, 0, lw, footer_h, fill=1, stroke=0)
    # Labels and title
    c.setFillColor(colors.black)
    c.setFont('Helvetica-Bold', 5.8)
    c.drawString(0.12 * inch, lh - 0.15 * inch, 'ARENA TERRITORY' if is_arena else 'TERRITORY')
    fit_landscape_title(c, territory['name'], lw / 2, lh - 0.31 * inch, lw - 0.36 * inch, 13.0, 7.2)
    # Main text
    body_x = 0.18 * inch
    body_y = footer_h + 0.12 * inch
    body_w = lw - 0.36 * inch
    body_h = lh - header_h - footer_h - 0.22 * inch
    fit_paragraph(c, safe_html(territory['text']), body_x, body_y, body_w, body_h, 8.2, 5.4, 1.12)
    # Footer metadata
    c.setFillColor(colors.black)
    c.setFont('Helvetica', 4.6)
    c.drawString(0.10 * inch, 0.075 * inch, (deck_name.upper() if deck_name else 'MASTER POOL'))
    c.drawCentredString(lw / 2, 0.075 * inch, COPYRIGHT_CARD)
    c.drawRightString(lw - 0.10 * inch, 0.075 * inch, f'{VERSION}{("  "+serial) if serial else ""}')
    c.restoreState()


def page_positions() -> List[Tuple[float, float]]:
    return src.page_positions()


def draw_card_pages(path: Path, items: List[Dict[str, Any]], title_meta: str) -> None:
    c = canvas.Canvas(str(path), pagesize=letter)
    positions = page_positions()
    sheets = math.ceil(len(items)/9)
    for pidx in range(sheets):
        page_items = items[pidx*9:(pidx+1)*9]
        for idx, item in enumerate(page_items):
            x, y = positions[idx]
            kind = item['kind']
            if kind == 'main':
                draw_main_card_with_copyright(c, x, y, item['card'], item.get('deck'), item.get('serial', ''))
            elif kind == 'territory':
                draw_territory_card_landscape(c, x, y, item['territory'], item.get('deck'), item.get('serial', ''))
            elif kind == 'heartland':
                draw_heartland_card_with_copyright(c, x, y, item.get('label', 'HEARTLAND'), item.get('deck'))
            elif kind == 'aid':
                draw_aid_card_with_copyright(c, x, y, item['title'], item['lines'], item.get('deck'))
        c.setFont('Helvetica', 5)
        c.setFillColor(colors.HexColor('#666666'))
        c.drawCentredString(PAGE_W/2, 0.07*inch, f'{title_meta} - sheet {pidx+1} of {sheets} - {COPYRIGHT_SHORT}')
        c.showPage()
    c.save()


def build_four_deck_pdf() -> Path:
    items = []
    for deck_name, deck in src.DECKS.items():
        seq = Counter()
        for name in deck['cards']:
            seq[name] += 1
            items.append({'kind': 'main', 'card': src.CARD_BY_NAME[name], 'deck': deck_name, 'serial': f'{seq[name]}'})
        for t in deck['territories']:
            items.append({'kind': 'territory', 'territory': src.TERRITORY_BY_NAME[t], 'deck': deck_name})
        items.append({'kind': 'heartland', 'deck': deck_name})
        items.append({'kind': 'aid', 'deck': deck_name, 'title': 'TURN SEQUENCE', 'lines': [
            'Capture check.', 'Draw one card.', 'Optional Action before movement.', 'Advance, hold, or withdraw.', 'Optional Action after movement if unused.', 'Cleanup to three cards.', 'Asset bank = Territories controlled.'
        ]})
        items.append({'kind': 'aid', 'deck': deck_name, 'title': 'BATTLE & DESTINATIONS', 'lines': [
            'Optional one card from hand.', 'Draw three; select up to one.', 'Reveal; cancel; resolve effects.', 'Roll and determine winner.', 'Loser retreats; attacker may occupy.', 'Hand to GY; battle-draw to discard. Heartland: ties/+1.'
        ]})
    assert len(items) == 144
    path = ROOT / 'Gauntlet_v0.5.6_Four_Deck_Playtest_Set.pdf'
    draw_card_pages(path, items, 'Gauntlet v0.5.6 Four-Deck Playtest Set')
    return path


def build_master_pool_pdf() -> Path:
    items = [{'kind': 'main', 'card': card} for card in src.CARDS]
    items += [{'kind': 'territory', 'territory': t} for t in src.TERRITORIES]
    items.append({'kind': 'heartland', 'label': 'HEARTLAND A'})
    items.append({'kind': 'heartland', 'label': 'HEARTLAND B'})
    assert len(items) == 81
    path = ROOT / 'Gauntlet_v0.5.6_Master_Card_Pool.pdf'
    draw_card_pages(path, items, 'Gauntlet v0.5.6 Master Card Pool')
    return path


def build_playtest_forms_pdf() -> Path:
    # Use source form builder, then stamp copyright with a simple footer overlay by regenerating from existing code pattern would be overkill.
    path = src.build_playtest_forms_pdf()
    # Rebuild in place with copyright footer to avoid external PDF overlay dependencies.
    c = canvas.Canvas(str(path), pagesize=letter)
    for page in range(4):
        margin=0.55*inch
        c.setFont('Helvetica-Bold',18); c.drawString(margin,PAGE_H-margin,'GAUNTLET PLAYTEST RECORD')
        c.setFont('Helvetica-Bold',10); c.drawRightString(PAGE_W-margin,PAGE_H-margin+2,VERSION)
        y=PAGE_H-margin-0.38*inch
        fields=[('Date / location',2.0),('Players',2.0),('Deck A / player',2.4),('Deck B / player',2.4),('First player',1.8),('Winner',1.8),('Ending',2.8),('Elapsed time',1.8),('Player-turns',1.8),('Battles',1.4),('Captures',1.4),('Reshuffles A / B',2.0),('Final Graveyard A / B',2.3)]
        colx=[margin,PAGE_W/2+0.12*inch]
        idx=0
        for row in range(7):
            for col in range(2):
                if idx>=len(fields): break
                label,_=fields[idx]; x=colx[col]
                c.setFont('Helvetica-Bold',7.5); c.drawString(x,y,label.upper())
                c.line(x,y-0.08*inch,x+(PAGE_W/2-margin-0.2*inch),y-0.08*inch)
                idx+=1
            y-=0.42*inch
        y-=0.05*inch
        c.setFont('Helvetica-Bold',9); c.drawString(margin,y,'PLAYER RATINGS (1 = poor, 5 = excellent)')
        y-=0.25*inch
        c.setFont('Helvetica',8)
        for label in ['Enjoyment','Agency','Pacing','Desire to replay']:
            c.drawString(margin,y,label+' - Player A:  1  2  3  4  5        Player B:  1  2  3  4  5')
            y-=0.25*inch
        prompts=['Rules questions or ambiguous interactions','Most valuable card or Territory','Least valuable card or Territory','Did either player feel unable to act? Why?','Was the winner apparent long before the game ended?','Other observations']
        for prompt in prompts:
            c.setFont('Helvetica-Bold',8); c.drawString(margin,y,prompt.upper()); y-=0.13*inch
            lines=2 if prompt!='Other observations' else 4
            for _ in range(lines):
                c.line(margin,y,PAGE_W-margin,y); y-=0.22*inch
            y-=0.07*inch
        c.setFont('Helvetica',5.5)
        c.drawCentredString(PAGE_W/2,0.28*inch,f'Gauntlet {VERSION} - Record sheet {page+1} of 4')
        c.drawCentredString(PAGE_W/2,0.18*inch,COPYRIGHT_SHORT)
        c.showPage()
    c.save()
    return path


def update_json(path: Path) -> Path:
    data = json.loads(path.read_text(encoding='utf-8'))
    data['asset_bank'] = {
        'terminology': 'Players bank Assets in their Asset bank.',
        'limit': "A player's Asset bank limit equals the number of Territories they control.",
        'occupied_territories': 'Occupied Territories do not count until captured.',
        'heartlands': 'Heartlands are not Territories and do not increase the Asset bank limit.',
        'limit_decrease': 'If the limit decreases below the number of Assets a player controls, they immediately choose and discard Assets until within the limit.'
    }
    data['action_destinations'] = {
        'banked_asset': "Action text that says to bank the card as an Asset places it in the player's Asset bank, subject to the Asset bank limit.",
        'condition': 'Action text with an ongoing or delayed effect that is not banked as an Asset, discarded after resolving, or placed as an Overlay or Territory is played as a Condition and remains until its card text directs it to be discarded.',
        'one_shot': 'One-shot Actions resolve and are discarded unless their text says otherwise.'
    }
    data['copyright'] = {
        'notice': 'Copyright © 2026 Tymon Scott. All rights reserved.',
        'owner': 'Tymon Scott',
        'use': 'Provided for private review, playtesting, and development reference only. No license is granted except personal printing for Gauntlet playtesting and feedback.',
        'feedback': 'Submitting feedback grants Tymon Scott permission to use, adapt, and incorporate that feedback into Gauntlet without compensation or obligation.'
    }
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding='utf-8')
    return path


def update_deck_lists(path: Path) -> Path:
    text = path.read_text(encoding='utf-8')
    notice = f'**Copyright:** {COPYRIGHT_SHORT}\n\n'
    if '**Copyright:**' not in text:
        text = text.replace(f'**Date:** {DATE}\n\n', f'**Date:** {DATE}\n\n{notice}')
    path.write_text(text, encoding='utf-8')
    return path


def write_copyright() -> Path:
    text = '''# Copyright and permitted use

Copyright © 2026 Tymon Scott. All rights reserved.

Gauntlet, including its rulebook text, card text, card lists, game materials, print-and-play files, graphic layout, documentation, and related design materials, is owned by Tymon Scott unless otherwise noted.

This repository is made available for private review, playtesting, and development reference only. No license is granted to copy, distribute, publish, sell, sublicense, create derivative commercial works from, or otherwise exploit these materials without written permission from the copyright owner.

You may download and print the included materials solely for personal playtesting and feedback related to Gauntlet.

Submitting comments, suggestions, issues, or playtest feedback does not transfer ownership of Gauntlet. By submitting feedback, you grant Tymon Scott permission to use, adapt, and incorporate that feedback into Gauntlet without compensation or obligation.

All rights reserved.
'''
    path = ROOT / 'COPYRIGHT.md'
    path.write_text(text, encoding='utf-8')
    return path


def write_contributing() -> Path:
    text = '''# Contributing and playtest feedback

Gauntlet is an unpublished playtest project owned by Tymon Scott.

Feedback, rules questions, playtest reports, bug reports, balance suggestions, and wording suggestions are welcome. By submitting comments, issues, pull requests, suggested card text, playtest notes, or other feedback, you grant Tymon Scott permission to use, adapt, modify, publish, and incorporate that feedback into Gauntlet without compensation, credit obligation, or transfer of ownership.

Do not submit material you do not have the right to share. Do not submit third-party copyrighted artwork, card text, game text, lore, or other protected material unless you have permission to do so.

Submitting feedback does not give you ownership of Gauntlet or any portion of the project. No license is granted to copy, redistribute, sell, publish, sublicense, or create derivative commercial works from Gauntlet materials.

For ordinary playtest reports, please include:

- Version number used
- Decks used
- Winner and ending type
- Approximate elapsed time
- Approximate player-turn count
- What felt fun
- What felt slow, confusing, or unfair
- Any card or Territory interaction that required a ruling
'''
    path = ROOT / 'CONTRIBUTING.md'
    path.write_text(text, encoding='utf-8')
    return path


def write_readme() -> Path:
    text = f'''# Gauntlet {VERSION} Print-and-Play Package

Generated {DATE}. Asset Bank Patch release with copyright notices and landscape-oriented Territory card text.

## Files

- `Gauntlet_v0.5.6_Complete_Playtest_Guide.docx` - editable canonical guide.
- `Gauntlet_v0.5.6_Complete_Playtest_Guide.pdf` - print-ready guide.
- `Gauntlet_v0.5.6_Four_Deck_Playtest_Set.pdf` - four complete 30-card decks, twelve Territories, four Heartlands, and eight player aids.
- `Gauntlet_v0.5.6_Master_Card_Pool.pdf` - one copy of all 54 main-deck cards, all 25 Territories, and two Heartlands.
- `Gauntlet_v0.5.6_Playtest_Forms.pdf` - four game record sheets.
- `Gauntlet_v0.5.6_Canonical_Data.json` - editable structured source data.
- `Gauntlet_v0.5.6_Playtest_Deck_Lists.md` - deck manifests.
- `COPYRIGHT.md` - copyright and permitted-use notice.
- `CONTRIBUTING.md` - feedback/contribution terms.

## Printing

Card sheets use exact 2.5 x 3.5 inch poker-card fronts in a 3 x 3 US Letter layout. Print at **Actual Size / 100%**, with scaling disabled. For a durable prototype, cut the fronts and sleeve them in opaque standard poker sleeves with ordinary playing cards or bulk trading cards behind them. This avoids duplex alignment problems.

Territory cards are still cut to standard poker-card size, but their text is printed in landscape orientation so the card can face the controlling player on the board.

The Four-Deck Playtest Set is organized deck-by-deck and labels every card at the bottom for sorting. The Shadow Directorate is the v0.5.6 Attrition trial list; Attrition replaces Court Martial at equal cost.

## Rules changes incorporated

- Heartland terminology.
- Hand-origin battle cards go to the Graveyard.
- Battle-drawn battle cards go to the discard pile.
- Replayed Battle cards normally go to the Graveyard.
- Incomplete draws draw as many cards as possible.
- Attrition uses the locked three-card Battle effect.
- Last Stand: Heartland defenders have Homeland Advantage and gain +1.
- Human playtest errata for Conscription, Watchtower, Fortifications, and Embargo.
- Asset bank terminology: players bank Assets in the Asset bank.
- Asset bank limit equals the number of Territories controlled.
- Action destination clarification: delayed or ongoing Actions that are not banked as Assets are played as Conditions and discarded when directed by card text.

## Copyright and use

Gauntlet is currently an unpublished playtest project.

Copyright © 2026 Tymon Scott. All rights reserved.

The materials in this repository are provided for private review and playtesting only. They may not be copied, redistributed, sold, republished, or used to create commercial derivative works without written permission.
'''
    path = ROOT / 'README.md'
    path.write_text(text, encoding='utf-8')
    return path


def write_release_notes() -> Path:
    text = f"""# Gauntlet {VERSION} - Asset Bank Patch

Release date: {DATE}

Copyright © 2026 Tymon Scott. All rights reserved. Playtest copy.

## Rules updates

- Renamed the persistent Asset area to the **Asset bank**.
- Standardized the verb **bank** for placing a card into the Asset bank as an Asset.
- Asset bank limit now equals the number of Territories a player controls.
- Occupied Territories do not count toward the Asset bank limit until captured.
- Heartlands are not Territories and do not increase the Asset bank limit.
- If a player's Asset bank limit decreases below their number of banked Assets, that player immediately chooses and discards Assets until within the limit.
- Clarified Action-card destinations: Actions that are not banked as Assets, discarded after resolving, or placed as Overlays/Territories are played as Conditions and discarded when directed by the card text.

## Card wording updates

- Updated Asset cards to use "Bank [card] as an Asset." wording.
- Updated delayed or ongoing Action cards to state when they are played as Conditions and when they are discarded.
- Updated Conscription to refer to cards whose Action banks them as Assets.
- Updated Scorched Earth and related effects to refer to banking Assets rather than playing Assets.

## Production/document updates

- Retained the v0.5.5 human-playtest errata for Conscription, Watchtower, Fortifications, and Embargo.
- Retained copyright notices, repository-facing copyright files, and landscape-oriented Territory cards.
- No Foothold, failed-counterattack capture rule, or other pacing-system rule is included in this patch.
"""
    path = ROOT / 'Gauntlet_v0.5.6_Release_Notes.md'
    path.write_text(text, encoding='utf-8')
    return path

def make_zip(paths: List[Path]) -> Path:
    zpath = ROOT / 'Gauntlet_v0.5.6_Print_and_Play_Package.zip'
    with zipfile.ZipFile(zpath, 'w', zipfile.ZIP_DEFLATED) as z:
        for p in paths:
            if p.exists():
                z.write(p, arcname=p.name)
    return zpath


def main():
    apply_v056_changes()
    src.validate()
    outputs: List[Path] = []
    outputs.append(src.export_json())
    outputs[-1] = update_json(outputs[-1])
    outputs.append(src.export_deck_lists_md())
    outputs[-1] = update_deck_lists(outputs[-1])
    outputs.append(src.build_docx())
    outputs[-1] = patch_docx(outputs[-1])
    outputs.append(build_four_deck_pdf())
    outputs.append(build_master_pool_pdf())
    outputs.append(build_playtest_forms_pdf())
    outputs.append(write_readme())
    outputs.append(write_copyright())
    outputs.append(write_contributing())
    outputs.append(write_release_notes())
    # Copy this build script into package as the source generator.
    outputs.append(Path(__file__))
    # PDF version of guide is generated externally by render_docx.py; include after build if present.
    print('\n'.join(str(p) for p in outputs))

if __name__ == '__main__':
    main()
