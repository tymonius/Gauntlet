from __future__ import annotations

import json
import math
import os
import re
import textwrap
import zipfile
from collections import Counter, defaultdict
from pathlib import Path
from typing import Dict, List, Any, Iterable, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, KeepInFrame

ROOT = Path('/mnt/data/gauntlet_v0_5_4')
VERSION = 'v0.5.4'
DATE = 'June 21, 2026'

# ---------------------------------------------------------------------------
# Canonical data
# ---------------------------------------------------------------------------

CARDS: List[Dict[str, Any]] = [
    {
        'name': 'Arcane Knowledge', 'cost': 5,
        'action': 'Place Arcane Knowledge in your Assets area. Once per turn during a battle, Arcane Knowledge may use the eligible Battle effect of one other card you played in that battle.',
        'battle': 'Arcane Knowledge may use the eligible Battle effect of one other card you played in this battle.',
        'reminder': 'It cannot use cancellation, special-reveal, battle-ending, follow-up-battle, source-dependent, or card-lifecycle effects.'
    },
    {
        'name': 'Armistice', 'cost': 4,
        'action': "Place Armistice in your Conditions area. For the rest of this turn and throughout your opponent's next turn, neither player may initiate a battle. Discard Armistice at the end of that turn.",
        'battle': "Resolve effects that cancel Battle cards first. If Armistice remains in play, end the battle immediately without a winner. The attacker returns to the Territory they entered from. Other unresolved Battle effects do not resolve. Place all other Battle cards still in play in their owners' discard piles, and place Armistice in your Graveyard.",
        'reminder': "The attacker's return is not a retreat. Effects that already resolved are not undone."
    },
    {
        'name': 'Assassins', 'cost': 4,
        'action': "Look at your opponent's hand. Choose one card and place it in their discard pile.",
        'battle': "Choose one opposing Battle card committed from hand. That card has no effect; place it in its owner's Graveyard immediately. If your opponent committed no card from hand, they gain disadvantage during this battle."
    },
    {
        'name': 'Assimilation', 'cost': 4,
        'action': 'Place Assimilation in your Conditions area. If you win your next battle this turn as the attacker on a Territory your opponent controls, immediately capture that Territory instead of occupying it. If another effect would delay capture, reduce that delay by one round instead. Discard Assimilation at the end of the turn.',
        'battle': 'If you win this battle as the attacker on a Territory your opponent controls, immediately capture that Territory instead of occupying it. This effect overrides effects that would delay capture.',
        'reminder': 'Against Protracted Siege, the Action restores the normal capture schedule; the Battle effect captures immediately.'
    },
    {
        'name': 'Attrition', 'cost': 3,
        'action': "Place Attrition in your Assets area. Whenever your opponent loses a battle against you, each card they played from their battle draw goes to their Graveyard instead of their discard pile.",
        'battle': "If your opponent loses this battle, place every card from their initial battle draw in their Graveyard.",
        'reminder': 'The initial battle draw is the first three cards drawn. This includes cards already discarded. Ignore additional and replacement cards; if fewer than three were drawn, affect all of them.'
    },
    {
        'name': 'Blockade', 'cost': 4,
        'action': "Place Blockade in your opponent's Conditions area. At the beginning of each of their turns, after their normal draw, if they have more than one card in hand, they discard one at random. When they win a battle, discard Blockade.",
        'battle': "Your opponent may discard one card from their hand. If they do not, each card they played from their battle draw has no effect during this battle. If they lose, place Blockade in their Conditions area instead of its normal destination.",
        'reminder': 'A player may have only one Blockade Condition.'
    },
    {
        'name': 'Brothers in Arms', 'cost': 4,
        'action': 'Place Brothers in Arms in your Assets area. Whenever you play both a card from your hand and a card from your battle draw in the same battle, gain advantage. If you win, place the hand-origin card in your discard pile instead of your Graveyard.',
        'battle': 'If Brothers in Arms was played from your hand and you also play a card from your battle draw - or vice versa - gain advantage. The other card cannot be canceled or have its Battle effect ignored.',
        'reminder': 'Brothers in Arms counts as one of the pair, but it can still be canceled.'
    },
    {
        'name': 'Capital Gains', 'cost': 2,
        'action': 'As an additional cost to play Capital Gains, discard one other card from your hand. Place Capital Gains in your Conditions area. At the beginning of your next turn, draw three cards instead of your normal draw, then discard Capital Gains.',
        'battle': 'After this battle, if you won, draw two cards. If you lost, discard one card from your hand.',
        'reminder': 'You cannot play the Action if you have no other card to discard.'
    },
    {
        'name': 'Capital Punishment', 'cost': 4,
        'action': "Choose one opposing Asset. If you defeat that opponent in battle this turn, place the chosen Asset in its owner's Graveyard.",
        'battle': "Choose one opposing Battle card. That card has no effect; place it in its owner's Graveyard immediately."
    },
    {
        'name': 'Conscription', 'cost': 3,
        'action': 'Draw one card. You may immediately play one card from your hand whose Action places it in your Assets area.',
        'battle': 'Reveal Conscription before the other Battle cards. You may commit one additional card from your hand face down.'
    },
    {
        'name': 'Contraband', 'cost': 3,
        'action': 'Choose one card in your discard pile and return it to your hand.',
        'battle': 'When Contraband is revealed, choose one card in your discard pile whose Battle effect can still resolve. Place Contraband in your Graveyard and play the chosen card in its place.',
        'reminder': 'The replayed card follows its own Battle text and normally enters the Graveyard after battle.'
    },
    {
        'name': 'Counterintelligence', 'cost': 2,
        'action': 'Place Counterintelligence in your Assets area. Opposing effects cannot look at, reveal, or require you to reveal your hand, battle draw, face-down Battle cards, or face-down Territories. This does not prevent revelation required by the normal rules of battle or Territory exploration.',
        'battle': 'Add +1 to your battle total. Until the normal reveal, opposing effects cannot look at or reveal your face-down Battle cards.'
    },
    {
        'name': 'Court Martial', 'cost': 3,
        'action': "Place Court Martial in your opponent's Conditions area. After their next battle, discard Court Martial. If they lost that battle, they must retreat one additional tile.",
        'battle': 'Your opponent gains disadvantage during this battle. If they lose, they must retreat one additional tile.',
        'reminder': 'The additional retreat is forced.'
    },
    {
        'name': 'Decoys', 'cost': 3,
        'action': 'Place Decoys in your Assets area. If an opposing effect would cause one or more of your other Assets to leave play, you may discard Decoys. If you do, choose one affected Asset; it remains in play.',
        'battle': 'Opposing effects that cancel Battle cards must target Decoys before they can target your other Battle cards.'
    },
    {
        'name': 'Embargo', 'cost': 2,
        'action': "Place Embargo in your opponent's Conditions area. At the beginning of their next turn, after their normal draw, they discard one card from their hand at random. Then discard Embargo.",
        'battle': "Choose one opposing Battle card. That card has no effect; return it to its owner's hand immediately."
    },
    {
        'name': 'Entrenchment', 'cost': 2,
        'action': 'Place Entrenchment in your Assets area. When your opponent advances onto a Territory adjacent to your token, their movement ends and they cannot play an Action card after movement that turn.',
        'battle': 'If you are the defender, your opponent gains disadvantage during this battle.'
    },
    {
        'name': 'Fealty', 'cost': 2,
        'action': 'Place Fealty in your Assets area. Opposing card effects cannot give you disadvantage.',
        'battle': 'Ignore one disadvantage affecting you during this battle. If you have no disadvantage, add +1 to your battle total instead.'
    },
    {
        'name': 'Fog of War', 'cost': 2,
        'action': 'Place Fog of War over one revealed Territory. During the next battle fought there, each player randomly selects from their battle draw instead of choosing. After that battle, discard Fog of War.',
        'battle': 'Your opponent randomly selects from their battle draw instead of choosing during this battle.',
        'reminder': 'Randomly select the number of battle-drawn cards that player is currently allowed to play.'
    },
    {
        'name': 'Fortifications', 'cost': 3,
        'action': 'Place Fortifications in your Assets area. When defending, you may play up to two cards from your battle draw instead of one.',
        'battle': 'If you are the defender, draw one additional battle card before choosing which battle-drawn card to play.',
        'reminder': 'The Battle effect increases selection only; an active Fortifications Asset sets the maximum at two battle-drawn cards.'
    },
    {
        'name': 'Illegal Occupation', 'cost': 3,
        'action': 'Place Illegal Occupation in your Assets area. While your opponent occupies a Territory you control without controlling it, their Assets are inactive.',
        'battle': 'If you are counterattacking an opponent occupying a Territory you control, their Assets are inactive during this battle and you gain advantage.'
    },
    {
        'name': 'Insurrection', 'cost': 4,
        'action': 'Discard your hand. Each player shuffles their discard pile into their deck. Draw three cards, then you may immediately play one additional Action card.',
        'battle': 'If you are the attacker, gain advantage. If you are counterattacking an opponent occupying a Territory you control, gain double advantage instead.'
    },
    {
        'name': 'Invasion', 'cost': 4,
        'action': 'Gain two additional movements this turn. Movements granted by Invasion may be used only to advance.',
        'battle': 'If you are the attacker, draw one additional battle card. You may play that card in addition to your other Battle cards.',
        'reminder': 'Unused movement is lost when a battle begins. If the additional Battle card is not played, discard it normally.'
    },
    {
        'name': 'Liberation', 'cost': 4,
        'action': 'Place Liberation in your Assets area. Whenever you win a counterattack against an opponent occupying a Territory you control, draw one card and you may immediately play one additional Action card.',
        'battle': 'If you win a counterattack against an opponent occupying a Territory you control, after that battle fully resolves, you may advance one tile. If this begins another battle, gain advantage during that battle.'
    },
    {
        'name': 'Manifest Destiny', 'cost': 5, 'unique': True,
        'action': 'Insert Manifest Destiny face down between your Heartland and the Territory nearest your Heartland. It becomes a new blank Territory you control.',
        'battle': 'If you win as the attacker on an enemy-controlled Territory, insert Manifest Destiny face down between the contested Territory and the space from which you attacked. It becomes a new blank Territory you control instead of going to its normal destination.',
        'reminder': 'Manifest Destiny is a permanent Territory, not an Overlay. Maximum one copy per deck.'
    },
    {
        'name': 'Militias', 'cost': 3,
        'action': 'Place Militias in your Assets area. During battles on a Territory you control, your opponent gains disadvantage.',
        'battle': 'Your opponent gains disadvantage. If this battle is on a Territory you control, they gain double disadvantage instead.'
    },
    {
        'name': 'Monetary Crisis', 'cost': 2,
        'action': 'Each player discards their hand, then draws two cards.',
        'battle': 'After this battle, each player discards down to one card.'
    },
    {
        'name': 'Necromancy', 'cost': 5,
        'action': 'Return one non-Necromancy card from your Graveyard to your hand. Place Necromancy in your Graveyard instead of your discard pile.',
        'battle': 'After this battle, return one non-Necromancy card from your Graveyard to your hand.'
    },
    {
        'name': 'New Recruits', 'cost': 1,
        'action': 'Discard one other card from your hand, then draw two cards.',
        'battle': 'Add +1 to your battle total.',
        'reminder': 'You cannot play the Action if you have no other card to discard.'
    },
    {
        'name': 'Palisade Wall', 'cost': 2,
        'action': "Place Palisade Wall in your Conditions area. During the next battle in which you are the defender, your opponent's Assets are inactive. Discard Palisade Wall after that battle.",
        'battle': 'If you are the defender, choose one opposing Battle card committed from hand. That card has no effect during this battle. If your opponent committed no card from hand, gain advantage instead.',
        'reminder': 'The chosen card remains played and goes to its normal destination after battle.'
    },
    {
        'name': 'Patriotism', 'cost': 3,
        'action': 'Place Patriotism in your Assets area. During battles on a Territory you control, the first +1 or advantage granted by one of your Battle cards is doubled.',
        'battle': 'If you are defending a Territory you control, gain advantage, and your Homeland Advantage cannot be disabled during this battle.',
        'reminder': '+1 becomes +2; advantage becomes double advantage. The Action applies only once per battle.'
    },
    {
        'name': 'Protracted Siege', 'cost': 4,
        'action': 'Place Protracted Siege in your Assets area. Whenever your opponent occupies a Territory you control, they do not capture it at the beginning of their next turn. If they still occupy it at the beginning of the following turn, they capture it normally.',
        'battle': 'If you are defending a Territory you control and lose, place Protracted Siege in your opponent\'s Conditions area instead of its normal destination. They do not capture the Territory occupied during this battle at the beginning of their next turn. If they still occupy it at the beginning of the following turn, they capture it normally, then discard Protracted Siege. If they cease occupying it first, discard Protracted Siege.'
    },
    {
        'name': 'Rallying Cry', 'cost': 1,
        'action': 'Draw one card.',
        'battle': 'Add +1 to your battle total.'
    },
    {
        'name': 'Redemption', 'cost': 2,
        'action': 'Place Redemption in your Conditions area. The next time an opposing effect causes one or more of your cards to enter your discard pile, choose one of those cards and return it to your hand after that effect resolves. Then discard Redemption.',
        'battle': 'If an opposing effect causes one other Battle card you played to have no effect and enter your discard pile, return that card to your hand after the battle instead.'
    },
    {
        'name': 'Reinforcements', 'cost': 2,
        'action': 'Place Reinforcements in your Conditions area. During your next turn, you may play one additional Action card. Discard Reinforcements at the end of that turn.',
        'battle': 'After the other Battle cards are revealed, draw one additional battle card. You may immediately play it face up if its Battle effect can still resolve.'
    },
    {
        'name': 'Resistance', 'cost': 3,
        'action': 'Place Resistance in your Assets area. When counterattacking an opponent occupying a Territory you control, draw two additional battle cards before choosing which battle-drawn card to play.',
        'battle': 'If you are counterattacking an opponent occupying a Territory you control, gain advantage. If you win, place Resistance in your Assets area instead of its normal destination.',
        'reminder': 'If your Assets area is full, you may discard one Asset to make room; otherwise Resistance goes to its normal destination.'
    },
    {
        'name': 'Revolution', 'cost': 4,
        'action': 'Each player discards their hand, then draws cards equal to the number of cards the other player discarded.',
        'battle': 'After all rerolls, you may exchange the two final selected die results. Each player keeps their own modifiers.',
        'reminder': 'If both players use Revolution, the exchanges cancel.'
    },
    {
        'name': 'Rousing Speech', 'cost': 2,
        'action': 'Place Rousing Speech in your Assets area. Whenever your opponent places a new card in their Assets area, you may draw one card, then discard one card.',
        'battle': 'If your opponent has more face-up Assets than you, gain advantage.',
        'reminder': 'Turning an existing Asset face up does not trigger Rousing Speech.'
    },
    {
        'name': 'Sabotage', 'cost': 2,
        'action': 'Choose one opposing Asset. Turn it face down until the start of your next turn.',
        'battle': "Choose one opposing Battle card. That card has no effect; place it in its owner's discard pile immediately."
    },
    {
        'name': 'Scorched Earth', 'cost': 4,
        'action': "Place Scorched Earth in your opponent's Conditions area. After their next battle, if they won, discard Scorched Earth. If they lost, they choose one Asset they control and discard it, and they cannot play Assets during their next turn. Discard Scorched Earth at the end of that turn.",
        'battle': 'If your opponent loses this battle, they choose one Asset they control and discard it. Place Scorched Earth in their Conditions area instead of its normal destination. They cannot play Assets during their next turn. Discard Scorched Earth at the end of that turn.'
    },
    {
        'name': 'Scouting Report', 'cost': 1,
        'action': "Look at the top card of your opponent's deck, one random card from their hand, or one face-down Territory.",
        'battle': 'Reveal Scouting Report before the other Battle cards. Look at one face-down opposing Battle card. You may replace Scouting Report with one unplayed card from your battle draw. If you do, place Scouting Report in your Graveyard immediately and place the replacement face down.'
    },
    {
        'name': 'Sedition', 'cost': 3,
        'action': 'Your opponent chooses one Asset they control and discards it.',
        'battle': 'Your opponent chooses one face-up Asset they control. It is inactive during this battle. If they have no face-up Assets, add +1 to your battle total.'
    },
    {
        'name': 'Shock and Awe', 'cost': 5,
        'action': 'Place Shock and Awe in your Conditions area. If you win your next battle this turn as the attacker on a Territory your opponent controls, immediately capture that Territory instead of occupying it, then you may advance one additional tile. If another effect would delay the capture, reduce that delay by one round instead; you may advance only if the Territory is captured. Discard Shock and Awe at the end of the turn.',
        'battle': 'If you win this battle as the attacker on a Territory your opponent controls, immediately capture that Territory instead of occupying it, then you may advance one additional tile. This capture overrides effects that would delay it. If the advance begins another battle, gain advantage during that battle.',
        'reminder': 'Shock and Awe can cause only one additional advance per turn.'
    },
    {
        'name': 'Siege Weaponry', 'cost': 5,
        'action': "Place Siege Weaponry face up over the nearest revealed enemy-controlled Territory ahead of you that has a printed effect. While it remains face up, that Territory's printed effect is inactive. The next time you attack there, discard Siege Weaponry if you lose. If you win, or if you capture that Territory without a battle, turn Siege Weaponry face down and leave it there.",
        'battle': "If you are attacking on an enemy-controlled Territory, place Siege Weaponry face up over it. That Territory's printed effect is inactive during this battle. If you win, turn Siege Weaponry face down and leave it there instead of placing it in its normal destination.",
        'reminder': 'Face-down Siege Weaponry is an active Ruins Overlay and remains until specifically repaired.'
    },
    {
        'name': 'Spies', 'cost': 1,
        'action': 'Place Spies in your Assets area. Your opponent keeps their hand face up.',
        'battle': 'Reveal Spies before the other Battle cards. Your opponent reveals every card committed from hand and their entire battle draw, then finalizes their battle-draw selection face up. You may then change your own battle-draw selection. If you replace Spies, place it in your Graveyard immediately.'
    },
    {
        'name': 'Stand Ground', 'cost': 2,
        'action': 'Until the start of your next turn, your token cannot be moved by opposing card effects.',
        'battle': 'If you are the defender, gain advantage.'
    },
    {
        'name': 'Strategic Withdrawal', 'cost': 3,
        'action': 'Return one Asset you control to your hand. If you do, gain one additional movement this turn.',
        'battle': 'If you lose this battle, after completing the required retreat, you may voluntarily withdraw one additional tile. If you do, return one other card you played during this battle to your hand instead of placing it in its normal destination.',
        'reminder': 'The additional withdrawal is voluntary and can trigger Refuge.'
    },
    {
        'name': 'Supplies', 'cost': 1,
        'action': 'Place Supplies in your Conditions area. At the beginning of your next turn, draw two additional cards, then discard Supplies.',
        'battle': 'After this battle, draw two cards, then discard one card.'
    },
    {
        'name': 'Tariffs', 'cost': 3,
        'action': 'Draw two cards. You may immediately play one additional Action card. Skip your normal draw at the beginning of your next turn.',
        'battle': 'Your opponent may discard one card from their hand. If they do not, add +1 to your battle total.',
        'reminder': 'Multiple Tariffs may chain, but the skipped-draw penalty does not stack.'
    },
    {
        'name': 'The Black Edict', 'cost': 4,
        'action': 'Each player chooses one Asset they control to keep and discards the rest.',
        'battle': 'All Assets are inactive during this battle.'
    },
    {
        'name': 'Treason', 'cost': 5,
        'action': "Place Treason in your opponent's Conditions area. During their next battle, after Battle cards are revealed, choose one eligible card they committed from hand. That card has no effect; place it in its owner's discard pile immediately. Treason uses that card's Battle effect for you. Then discard Treason.",
        'battle': "Choose one eligible opposing non-cancellation Battle card. That card has no effect; place it in its owner's discard pile immediately. Treason uses that card's Battle effect for you."
    },
    {
        'name': 'Tyranny', 'cost': 4,
        'action': "Place Tyranny in your Assets area. Once per battle, the first opposing Battle card that would add to your opponent's battle total or grant them advantage has those portions of its effect ignored.",
        'battle': "During this battle, opposing Battle cards cannot add to your opponent's battle total or grant them advantage.",
        'reminder': 'Other portions of an affected card still resolve.'
    },
    {
        'name': 'Valor', 'cost': 2,
        'action': 'Place Valor in your Assets area. Whenever you lose a battle, draw one card after retreating.',
        'battle': "After the battle dice are rolled, if your total is lower than your opponent's, you may reroll your die. You must use the new result."
    },
    {
        'name': 'War Crimes', 'cost': 4,
        'action': 'Place War Crimes in your Assets area. Whenever your opponent loses a battle against you, they cannot benefit from effects they control that trigger because they lost or retreated. Cards they played during that battle must go to their Graveyard unless an opposing effect moves them elsewhere.',
        'battle': 'If your opponent loses this battle, they cannot benefit from effects they control that trigger because they lost or retreated. Cards they played during this battle must go to their Graveyard unless an opposing effect moves them elsewhere. They must retreat one additional tile.',
        'reminder': 'The additional retreat is forced and does not trigger Refuge.'
    },
    {
        'name': 'Witchcraft', 'cost': 3,
        'action': 'Return one other card from your Graveyard to your discard pile.',
        'battle': "When Witchcraft is revealed, choose one non-Witchcraft card in your Graveyard whose eligible Battle effect can still resolve. Witchcraft uses that Battle effect for this battle. Leave the chosen card in your Graveyard."
    },
]

TERRITORIES: List[Dict[str, str]] = [
    {'name': 'Quicksand', 'text': 'If a player begins their Movement phase on this Territory, they cannot voluntarily withdraw or move more than one tile that turn. Forced movement is unaffected.'},
    {'name': 'Poisonous Gas', 'text': 'During a battle on this Territory, each player may play no more than one card and must place one card from their battle draw in their Graveyard. A battle-draw card played during that battle satisfies this requirement.'},
    {'name': 'Difficult Terrain', 'text': 'When a player enters this Territory, their movement ends for the turn. A player who starts their turn here or enters it during their turn cannot play an Action card after movement that turn.'},
    {'name': 'Garrison', 'text': 'When defending this Territory, its controller draws four battle cards instead of three.'},
    {'name': 'Disrupted Supply Lines', 'text': 'While a player is on this Territory, they may have no more than one face-up Asset. Turn their other Assets face down and inactive until they leave.'},
    {'name': 'Ruined Storehouse', 'text': 'When drawing at the beginning of their turn, a player on this Territory may draw the top card of their discard pile instead of their deck.'},
    {'name': 'Field Hospital', 'text': 'After a battle on this Territory, its controller may place one card they played that would enter their Graveyard in their discard pile instead.'},
    {'name': 'Exposed Flank', 'text': "When this Territory's controller counterattacks an opponent occupying it, that opponent cannot commit a card from hand during the battle."},
    {'name': 'High Ground', 'text': 'The defender in a battle on this Territory gains advantage.'},
    {'name': 'Fortified Pass', 'text': "When this Territory's controller defends it, the attacker's Assets are inactive for that battle."},
    {'name': 'Insurgency', 'text': 'While an opponent occupies this Territory without controlling it, their Assets are inactive.'},
    {'name': 'Watchtower', 'text': 'When the defender controls this Territory, any card the attacker commits from hand is placed face up instead of face down.'},
    {'name': 'Supply Depot', 'text': 'At the beginning of their turn, a player occupying and controlling this Territory draws two cards instead of one.'},
    {'name': 'Old Battlefield', 'text': 'After a battle on this Territory, its controller may place one unplayed card from their battle draw in their Graveyard instead of their discard pile.'},
    {'name': 'Refuge', 'text': 'When a player voluntarily withdraws onto this Territory, they draw one card.'},
    {'name': 'Command Tent', 'text': 'If you start your turn occupying and controlling this Territory, you may play one Action card before movement and one Action card after movement. If Command Tent is face down, you may reveal it at the start of your turn without using either Action opportunity.'},
    {'name': 'Monastery', 'text': "While this Territory's controller occupies it, cards cannot leave either player's Graveyard. During battles here, Witchcraft, Necromancy, and Arcane Knowledge have no effect."},
    {'name': 'Training Grounds', 'text': 'When defending this Territory, its controller may discard their entire battle draw and draw the same number of replacement cards. They must use the replacement draw.'},
    {'name': "King's Road", 'text': "When a player starts their turn on this Territory, they have two movements during that turn's Movement phase instead of one."},
    {'name': 'Toll Bridge', 'text': 'To advance from this Territory, a player must discard one card from their hand.'},
    {'name': "Smuggler's Pass", 'text': 'While occupying and controlling this Territory, a player may, instead of playing an Action card, place one card from their hand face down next to it. The stored card does not count toward their hand limit. At the beginning of a later turn while they control this Territory, they may return it to their hand. Only one card may be stored here. If control changes, discard it.'},
    {'name': 'Arena: Spoils of War', 'text': 'Homeland Advantage does not apply, and tied battle rolls are rerolled. After the battle, the winner may add one unplayed card from their battle draw to their hand instead of discarding it.'},
    {'name': 'Arena: No Quarter', 'text': 'Homeland Advantage does not apply, and tied battle rolls are rerolled. The loser retreats two tiles instead of one.'},
    {'name': 'Arena: Single Combat', 'text': "Homeland Advantage does not apply, and tied battle rolls are rerolled. Neither player's Assets apply during the battle."},
    {'name': 'Arena: Grand Melee', 'text': 'Homeland Advantage does not apply, and tied battle rolls are rerolled. Each player draws four battle cards instead of three and may play up to two of them.'},
]

DECKS: Dict[str, Dict[str, Any]] = {
    'Vanguard Coalition': {
        'subtitle': 'Flexible combined-arms offense',
        'cards': [
            'Rallying Cry','Rallying Cry','Rallying Cry','Redemption',
            'Scouting Report','Scouting Report','Scouting Report',
            'New Recruits','New Recruits','New Recruits','New Recruits',
            'Supplies','Stand Ground','Stand Ground','Entrenchment','Fealty',
            'Sabotage','Sabotage','Embargo','Counterintelligence','Conscription',
            'Fortifications','Strategic Withdrawal','Valor','Invasion','Assimilation',
            'Shock and Awe','Court Martial','Reinforcements','Rousing Speech'
        ],
        'territories': ["King's Road", 'High Ground', 'Supply Depot'],
        'status': 'v0.5.4 full-budget print list: one Rallying Cry is upgraded to Redemption; otherwise based on the previously simulation-tested list.'
    },
    'Bastion Coalition': {
        'subtitle': 'Defensive control with a finisher',
        'cards': [
            'Rallying Cry','Rallying Cry','Rallying Cry','Palisade Wall',
            'Scouting Report','Scouting Report','Scouting Report',
            'New Recruits','New Recruits','New Recruits','New Recruits',
            'Stand Ground','Stand Ground','Entrenchment','Palisade Wall','Fortifications',
            'Fealty','Sabotage','Sabotage','Shock and Awe','Court Martial',
            'Counterintelligence','Embargo','Strategic Withdrawal','Redemption','Valor',
            'Assimilation','Invasion','Rousing Speech','Reinforcements'
        ],
        'territories': ['High Ground','Command Tent','Ruined Storehouse'],
        'status': 'v0.5.4 full-budget print list: a second Palisade Wall replaces one Rallying Cry; otherwise based on Bastion Coalition v2.'
    },
    'Reclamation Front': {
        'subtitle': 'Counterattack and territorial recovery',
        'cards': [
            'Rallying Cry','Rallying Cry','Rallying Cry','Counterintelligence',
            'Scouting Report','Scouting Report','Scouting Report',
            'New Recruits','New Recruits','New Recruits','New Recruits',
            'Supplies','Stand Ground','Stand Ground','Entrenchment','Resistance',
            'Liberation','Illegal Occupation','Patriotism','Strategic Withdrawal','Redemption',
            'Valor','Sabotage','Embargo','Fealty','Assimilation','Invasion',
            'Court Martial','Reinforcements','Rousing Speech'
        ],
        'territories': ['Exposed Flank','Field Hospital',"King's Road"],
        'status': 'v0.5.4 full-budget print list: Counterintelligence replaces one Rallying Cry; otherwise based on the previously simulation-tested list.'
    },
    'Shadow Directorate': {
        'subtitle': 'Information warfare and attrition',
        'cards': [
            'Rallying Cry','Rallying Cry','Rallying Cry','Rallying Cry',
            'Scouting Report','Scouting Report','Scouting Report',
            'New Recruits','New Recruits','New Recruits','New Recruits',
            'Sabotage','Sabotage','Embargo','Assassins','Monetary Crisis','Sedition',
            'Counterintelligence','Stand Ground','Stand Ground','Entrenchment','Fealty',
            'Valor','Strategic Withdrawal','Conscription','Fortifications','Rousing Speech',
            'Assimilation','Invasion','Attrition'
        ],
        'territories': ['Watchtower','Supply Depot','Arena: Spoils of War'],
        'status': 'v0.5.4 trial list: Attrition replaces Court Martial at equal cost.'
    },
}

VERSION_HISTORY = [
    {'version':'v0.4.0','name':'Legacy Ruleset','changes':['Last coherent pre-overhaul ruleset.','Movement rolls and the earlier breaching/capture framework remained in use.']},
    {'version':'v0.5.0','name':'Core Rules Overhaul','changes':['Removed movement rolls.','Added advance, hold, and voluntary withdrawal.','Added occupation and one-turn counterattack capture window.','Rebuilt the battle sequence around a hand commitment plus a three-card battle draw.','Introduced Assets, Conditions, Overlays, and the current Territory framework.']},
    {'version':'v0.5.1','name':'Initial Balance Patch','changes':['Adjusted provisional card costs after initial simulation.','Introduced recommended well-rounded playtest decks.']},
    {'version':'v0.5.2','name':'Deck Flow Patch','changes':['Cards committed from hand during battle go to the Graveyard.','Cards played from the battle draw normally go to the discard pile.','Added the incomplete-draw rule.']},
    {'version':'v0.5.3','name':'Attrition Patch','changes':['Reworked Attrition around changing the destination of an opponent\'s battle-draw cards after a loss.','Attrition\'s Battle effect now sends the opponent\'s full initial three-card battle draw to the Graveyard after a loss.','Standardized Heartland terminology and synchronized card wording with the v0.5.2 lifecycle.']},
    {'version':'v0.5.4','name':'Last Stand Patch','changes':['Added Last Stand: a player defending in their own Heartland has Homeland Advantage and gains +1 to their battle total.','Clarified that Heartlands are not Territories and cannot be controlled or receive Territory effects unless a card explicitly says otherwise.','Updated Heartland cards, player aids, and the winning-game rules section.']},
]

# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------

CARD_BY_NAME = {c['name']: c for c in CARDS}
TERRITORY_BY_NAME = {t['name']: t for t in TERRITORIES}


def deck_cost(cards: Iterable[str]) -> int:
    return sum(CARD_BY_NAME[name]['cost'] for name in cards)


def validate() -> None:
    assert len(CARDS) == 54, len(CARDS)
    assert len(CARD_BY_NAME) == 54
    assert len(TERRITORIES) == 25, len(TERRITORIES)
    costs = Counter(c['cost'] for c in CARDS)
    assert costs == Counter({1:5,2:14,3:14,4:15,5:6}), costs
    for dname, deck in DECKS.items():
        assert len(deck['cards']) == 30, (dname, len(deck['cards']))
        missing = [n for n in deck['cards'] if n not in CARD_BY_NAME]
        assert not missing, (dname, missing)
        assert deck_cost(deck['cards']) <= 60, (dname, deck_cost(deck['cards']))
        assert len(deck['territories']) == 3
        assert len(set(deck['territories'])) == 3
        assert all(t in TERRITORY_BY_NAME for t in deck['territories'])
        assert sum(t.startswith('Arena:') for t in deck['territories']) <= 1
        assert deck['cards'].count('Manifest Destiny') <= 1

# ---------------------------------------------------------------------------
# Data exports
# ---------------------------------------------------------------------------


def export_json() -> Path:
    data = {
        'version': VERSION,
        'date': DATE,
        'status': 'Pre-release playtest edition',
        'deck_construction': {
            'minimum_cards': 30,
            'maximum_points': 60,
            'opening_hand': 3,
            'hand_limit': 3,
            'territories': 3,
            'maximum_arenas': 1,
            'unique_cards': ['Manifest Destiny'],
        },
        'battle_card_lifecycle': {
            'hand_origin': 'Graveyard',
            'battle_draw_origin': 'discard pile',
            'replayed': 'Graveyard unless an effect says otherwise',
            'unused_battle_draw': 'discard pile',
            'incomplete_draw': 'Draw as many as possible after reshuffling the discard pile.'
        },
        'heartland_last_stand': {
            'homeland_advantage': True,
            'battle_total_bonus': 1,
            'rule': 'When defending in your own Heartland, you have Homeland Advantage and gain +1 to your battle total.'
        },
        'version_history': VERSION_HISTORY,
        'cards': CARDS,
        'territories': TERRITORIES,
        'recommended_decks': {
            name: {**deck, 'card_count': len(deck['cards']), 'point_total': deck_cost(deck['cards'])}
            for name, deck in DECKS.items()
        }
    }
    path = ROOT / 'Gauntlet_v0.5.4_Canonical_Data.json'
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding='utf-8')
    return path


def export_deck_lists_md() -> Path:
    lines = [
        '# Gauntlet v0.5.4 Playtest Deck Lists', '',
        f'**Date:** {DATE}', '',
        'Each deck contains 30 cards and no more than 60 points. Territories are listed in order from the deck owner\'s Heartland toward the center.', ''
    ]
    for name, deck in DECKS.items():
        counts = Counter(deck['cards'])
        lines += [f'## {name}', '', f'*{deck["subtitle"]}*', '', f'**Cards:** 30  ', f'**Points:** {deck_cost(deck["cards"])}  ', f'**Territories:** {", ".join(deck["territories"])}  ', f'**Status:** {deck["status"]}', '', '| Qty | Card | Cost each |', '|---:|---|---:|']
        for card_name in sorted(counts):
            lines.append(f'| {counts[card_name]} | {card_name} | {CARD_BY_NAME[card_name]["cost"]} |')
        lines.append('')
    path = ROOT / 'Gauntlet_v0.5.4_Playtest_Deck_Lists.md'
    path.write_text('\n'.join(lines), encoding='utf-8')
    return path

# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def set_row_cant_split(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('Page ')
    run.font.size = Pt(8)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r = OxmlElement('w:r')
    r.append(fldChar1); r.append(instrText); r.append(fldChar2)
    paragraph._p.append(r)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)


def add_rule_box(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0,0)
    set_cell_shading(cell, 'E8E8E8')
    set_cell_margins(cell, 100, 140, 100, 140)
    p = cell.paragraphs[0]
    r = p.add_run(title.upper() + '\n')
    r.bold = True; r.font.size = Pt(9)
    r2 = p.add_run(text)
    r2.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_card_reference_cell(cell, card: Dict[str, Any]) -> None:
    set_cell_margins(cell, 85, 110, 85, 110)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(card['name'])
    r.bold = True; r.font.size = Pt(10.5)
    r2 = p.add_run(f'  COST {card["cost"]}' + (' | UNIQUE' if card.get('unique') else ''))
    r2.bold = True; r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(80,80,80)
    for label, key in [('ACTION','action'),('BATTLE','battle')]:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rr = p.add_run(label + ': '); rr.bold = True; rr.font.size = Pt(8.4)
        rr2 = p.add_run(card[key]); rr2.font.size = Pt(8.4)
    if card.get('reminder'):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run('REMINDER: '); rr.bold = True; rr.italic = True; rr.font.size = Pt(7.5)
        rr2 = p.add_run(card['reminder']); rr2.italic = True; rr2.font.size = Pt(7.5)


def build_docx() -> Path:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55); sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.62); sec.right_margin = Inches(0.62)

    styles = doc.styles
    styles['Normal'].font.name = 'Aptos'
    styles['Normal'].font.size = Pt(9.2)
    styles['Normal'].paragraph_format.space_after = Pt(4)
    for sty_name, size, color in [('Title',30,'111111'),('Subtitle',14,'555555'),('Heading 1',18,'111111'),('Heading 2',13,'333333'),('Heading 3',10.5,'555555')]:
        st = styles[sty_name]
        st.font.name = 'Aptos Display' if sty_name in ('Title','Heading 1','Heading 2') else 'Aptos'
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
    styles['Heading 1'].paragraph_format.space_before = Pt(10)
    styles['Heading 1'].paragraph_format.space_after = Pt(5)
    styles['Heading 2'].paragraph_format.space_before = Pt(7)
    styles['Heading 2'].paragraph_format.space_after = Pt(3)

    header = sec.header.paragraphs[0]
    header.text = f'GAUNTLET - COMPLETE PLAYTEST EDITION     {VERSION}'
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in header.runs:
        r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = RGBColor(90,90,90)
    add_page_number(sec.footer.paragraphs[0])

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run('GAUNTLET')
    r.bold = True; r.font.size = Pt(42); r.font.name = 'Aptos Display'
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run('Complete Playtest Edition'); r.font.size = Pt(20); r.bold=True
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run(VERSION); r.font.size = Pt(16); r.bold=True
    p4 = doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run(f'Rules, card reference, Territories, recommended decks, and playtest forms\n{DATE}').font.size = Pt(11)
    doc.add_paragraph('\n')
    add_rule_box(doc, 'Playtest status', 'This is a pre-release ruleset. Card costs and individual effects remain provisional, but all games recorded as v0.5.4 should use the rules and card text in this document without substitution.')
    doc.add_page_break()

    # Contents
    doc.add_heading('Contents', level=1)
    contents = [
        '1. Release identification and version history',
        '2. Game overview and components',
        '3. Deck construction and setup',
        '4. Turn, movement, and capture',
        '5. Battles',
        '6. Card zones and persistent effects',
        '7. Special rules and winning',
        '8. Territory reference',
        '9. Canonical card reference',
        '10. Recommended playtest decks',
        '11. Playtest procedure and record sheet',
        'Appendix A. Quick reference',
        'Appendix B. v0.5.4 change summary',
    ]
    for x in contents: add_bullet(doc, x)
    doc.add_heading('Using this package', level=2)
    add_bullet(doc, 'Use the Complete Playtest Guide as the canonical rules and wording source.')
    add_bullet(doc, 'Use the Four-Deck Print Set for structured matchup testing.')
    add_bullet(doc, 'Use the Master Card Pool to substitute or inspect individual cards and Territories.')
    add_bullet(doc, 'Record the exact version, deck names, elapsed time, player-turn count, and any rules questions after each game.')
    doc.add_page_break()

    # 1 Version history
    doc.add_heading('1. Release identification and version history', level=1)
    add_rule_box(doc, 'Current release', f'{VERSION} - Last Stand Patch. Released for physical playtesting on {DATE}.')
    for item in VERSION_HISTORY:
        doc.add_heading(f'{item["version"]} - {item["name"]}', level=2)
        for ch in item['changes']: add_bullet(doc, ch)
    doc.add_heading('Versioning convention', level=2)
    add_bullet(doc, '0.x.y indicates pre-release development.')
    add_bullet(doc, 'Increment the minor number for a new rules generation that makes prior playtest findings substantially obsolete.')
    add_bullet(doc, 'Increment the patch number for balance changes, card rewrites, timing corrections, and smaller rules refinements within the same generation.')
    add_bullet(doc, 'Formatting-only changes may retain the game version and use a revised document date.')

    # 2 Overview/components
    doc.add_heading('2. Game overview and components', level=1)
    doc.add_paragraph('Gauntlet is a two-player tactical card game about advancing across a shared line of Territories, defeating the opposing force, capturing ground, and reaching the enemy Heartland.')
    add_rule_box(doc, 'Objective', 'Advance through the opposing Territories and enter the enemy Heartland. If the opposing token occupies the Heartland, defeat it there to win. The Heartland defender makes a Last Stand: they have Homeland Advantage and gain +1 to their battle total.')
    doc.add_heading('Components', level=2)
    for x in [
        'One main deck for each player.', 'Three Territory cards for each player.',
        'One Heartland space or Heartland card for each player.', 'One token and at least one six-sided die for each player.',
        'A discard pile, face-up Graveyard, Assets area with three spaces, and Conditions area for each player.',
        'Space beside Territories for Overlays and stored cards.',
    ]: add_bullet(doc,x)

    # 3 deck/setup
    doc.add_heading('3. Deck construction and setup', level=1)
    doc.add_heading('Deck construction', level=2)
    for x in [
        'Your main deck must contain at least 30 cards and cost no more than 60 points.',
        'There is no general copy limit. Manifest Destiny is unique: maximum one copy per deck.',
        'Choose exactly three different Territory cards. Territories do not count toward the card minimum or point limit.',
        'You may choose no more than one Arena Territory.',
    ]: add_bullet(doc,x)
    doc.add_heading('Setup', level=2)
    for x in [
        'Each player secretly arranges their three Territories face down in a line extending from their Heartland toward the center. The two lines meet to form the Gauntlet.',
        'Players may inspect their own face-down Territories. The opponent may not.',
        'Place each token on its Heartland, just beyond its nearest Territory.',
        'Shuffle the main decks. Each player draws three cards.',
        'Roll a die to determine the first player; reroll ties.',
    ]: add_number(doc,x)
    add_rule_box(doc, 'Board order', 'Heartland A - three Territories selected by A - three Territories selected by B - Heartland B.')

    # key terms
    doc.add_heading('Key terms', level=2)
    key_terms = [
        ('CONTROL','A Territory is controlled by the player it is oriented toward. Rotate a captured Territory to face its new controller.'),
        ('ON A TERRITORY','A token is physically located on that Territory, regardless of who controls it.'),
        ('OCCUPATION','A player is occupying an opponent-controlled Territory after winning there but before capturing it.'),
        ('COUNTERATTACK','The controller attacks an enemy occupier during the capture window.'),
        ('ATTACKER','The player whose movement began the battle, including a player counterattacking an occupier.'),
        ('DEFENDER','The player whose token was already on the contested space when battle began.'),
        ('HOMELAND ADVANTAGE','A defender on a Territory they control, or a player defending in their own Heartland, wins tied battle totals.'),
        ('LAST STAND','When defending in your own Heartland, you have Homeland Advantage and gain +1 to your battle total.'),
        ('VOLUNTARY WITHDRAWAL','A player chooses to move one tile toward their own Heartland during Movement.'),
        ('FORCED RETREAT','A loser moves toward their own Heartland because of a battle result or another effect.'),
        ('INITIAL BATTLE DRAW','The first three cards a player attempts to draw for a battle before any replacement or additional battle cards.'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    set_repeat_table_header(table.rows[0])
    table.cell(0,0).text='TERM'; table.cell(0,1).text='MEANING'
    set_cell_shading(table.cell(0,0),'D9D9D9'); set_cell_shading(table.cell(0,1),'D9D9D9')
    for term,meaning in key_terms:
        row=table.add_row(); set_row_cant_split(row)
        row.cells[0].text=term; row.cells[1].text=meaning
        row.cells[0].paragraphs[0].runs[0].bold=True
        for c in row.cells: set_cell_margins(c)

    # 4 turn movement capture
    doc.add_heading('4. Turn, movement, and capture', level=1)
    doc.add_heading('Turn sequence', level=2)
    turn_steps = [
        'Capture check: If you still occupy an opponent-controlled Territory seized on your previous turn, capture it unless an effect delays capture.',
        'Draw: Draw one card. If the deck is empty, reshuffle the discard pile first. If the deck and discard pile cannot provide a card, draw nothing.',
        'First Action window: You may play one Action card.',
        'Movement: Advance one tile, hold position, or voluntarily withdraw one tile. Additional movement effects may grant more movement.',
        'Second Action window: If you did not play an Action card before movement, you may play one now.',
        'Cleanup: Discard down to the three-card hand limit.',
    ]
    for x in turn_steps: add_number(doc,x)
    add_rule_box(doc, 'Additional Actions', 'When a card allows an additional Action card, that card resolves normally and may grant further Action plays. There is no fixed chain limit unless a card says otherwise.')
    doc.add_heading('Movement', level=2)
    for x in [
        'Base movement requires no die roll.',
        'You may use fewer movements than you have available.',
        'With multiple movements, you may withdraw and later advance during the same turn.',
        'Entering an occupied space begins a battle and ends any unused movement.',
        'Tokens cannot move through or past one another.',
        'Forced retreat is not movement and does not consume movement opportunities.',
        'A token cannot retreat beyond its Heartland; stop at the Heartland if the full distance is unavailable.',
    ]: add_bullet(doc,x)
    doc.add_heading('Territory reveal, control, and capture', level=2)
    for x in [
        'Reveal a Territory when an opponent enters it or when its controller uses or triggers its effect.',
        'If entry begins a battle, reveal the Territory before either player selects Battle cards.',
        'Once revealed, a Territory remains face up.',
        'A face-down Territory is inert. Revealing it does not retroactively apply an entry effect to the movement that revealed it.',
        'When you could use a face-down Territory effect, you may reveal it instead of playing an Action card. Its effect applies immediately.',
        'If an attacker wins on an opponent-controlled Territory, the defender retreats and the attacker occupies it.',
        "The original controller has their next turn to counterattack. If the occupier is still there at the start of the occupier's following turn, capture it and rotate it toward the new controller.",
        'A counterattacker who wins retakes the space immediately because they already control it; no capture delay applies.',
    ]: add_bullet(doc,x)

    # 5 Battles
    doc.add_heading('5. Battles', level=1)
    doc.add_heading('Battle setup', level=2)
    battle_setup = [
        'Reveal the contested Territory if required and apply any effect that changes battle setup.',
        'Each player may commit up to one eligible Battle card from hand, face down, in their Hand row.',
        'Each player draws three cards as their initial battle draw. If fewer than three cards are available after reshuffling, draw as many as possible.',
        'Each player chooses up to one eligible card from their battle draw and places it face down in the Battle Draw row. Discard the unused cards unless an effect changes their destination.',
        'Resolve cards that say to reveal before the other Battle cards. If both players use such effects, resolve the attacker first, then the defender.',
        'Reveal the remaining Battle cards.',
    ]
    for x in battle_setup: add_number(doc,x)
    add_rule_box(doc, 'Card source', 'A card remains hand-origin, battle-drawn, or replayed for the entire battle. Replacing a card does not transfer the removed card\'s source to the replacement.')
    doc.add_heading('Battle-card effects', level=2)
    for x in [
        'Effects that cancel Battle cards resolve before other Battle effects. Within the same timing window, the attacker resolves first, then the defender.',
        'A canceled card has no effect and moves immediately to the destination named by the canceling card.',
        'A card whose effect is merely ignored remains played and moves to its normal destination after battle.',
        'Cards moved immediately by cancellation are no longer in the battle and may be affected in their new zone.',
        'Other effects resolve attacker first, then defender, when they share the same timing.',
        'Resolve as much of an effect as possible if one part is prohibited or impossible.',
    ]: add_bullet(doc,x)
    doc.add_heading('Battle dice', level=2)
    for x in [
        'Base battle: each player rolls one d6. Higher total wins.',
        'Advantage means roll one additional die and use the highest. Double advantage means roll two additional dice and use the highest.',
        'Disadvantage means roll one additional die and use the lowest. Double disadvantage means roll two additional dice and use the lowest.',
        'Advantage and disadvantage cancel one-for-one before rolling.',
        'Both players roll simultaneously.',
        'After the initial results are visible, optional rerolls resolve attacker first, then defender. A rerolled result must be used.',
        'After all rerolls, Revolution may exchange the two final selected natural die results. Modifiers remain with their original players.',
        'Apply +1 and other total modifiers after selecting the final dice.',
        'A defender on a Territory they control, or a player defending in their own Heartland, wins ties. All other ties are rerolled as part of the same battle; ongoing Battle effects remain active.',
    ]: add_bullet(doc,x)
    doc.add_heading('Resolving a battle', level=2)
    resolve_steps = [
        'Determine the winner and loser.',
        'Resolve effects triggered by winning or losing, including effects that change retreat or card destinations.',
        'The loser completes the normal forced retreat, then any additional forced retreats. Optional voluntary withdrawal occurs last.',
        'Establish the final position: a victorious attacker occupies the contested space; a victorious defender remains there.',
        'Resolve remaining after-battle effects, including draws, recovery, and cards that become Assets or Conditions.',
        'Move each remaining hand-origin Battle card to its owner\'s Graveyard, each remaining battle-drawn card to its owner\'s discard pile, and each replayed card to its owner\'s Graveyard, unless another effect changes its destination.',
        'Resolve at most one immediate post-battle advance from that victory. If multiple effects offer one, choose which effect to use.',
    ]
    for x in resolve_steps: add_number(doc,x)
    doc.add_heading('Follow-up battles', level=2)
    for x in [
        'The first battle must fully resolve before an additional advance.',
        'All one-battle effects expire and all first-battle cards move to their destinations.',
        'If the advance enters an occupied space, begin a completely new battle.',
        'Both players may commit new hand cards and receive a fresh battle draw.',
        'Assets, Conditions, Overlays, the new Territory, Homeland Advantage, and Last Stand are checked again.',
        'Only an effect that explicitly applies to the follow-up battle carries over.',
    ]: add_bullet(doc,x)

    # 6 zones
    doc.add_heading('6. Card zones and persistent effects', level=1)
    doc.add_heading('Default destinations', level=2)
    dest_table = doc.add_table(rows=1, cols=3); dest_table.style='Table Grid'; dest_table.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=dest_table.rows[0]; set_repeat_table_header(hdr)
    for i,h in enumerate(['CARD OR EVENT','DEFAULT DESTINATION','NOTES']):
        hdr.cells[i].text=h; set_cell_shading(hdr.cells[i],'D9D9D9')
    rows=[
        ('Resolved Action card','Discard pile','Unless it becomes an Asset, Condition, Overlay, Territory, or says otherwise.'),
        ('Unused battle-draw card','Discard pile','Includes unselected cards from the initial battle draw and later additional draws.'),
        ('Hand-origin card played in battle','Graveyard','This is the permanent cost of prepared commitment.'),
        ('Battle-drawn card played in battle','Discard pile','Attrition, War Crimes, Poisonous Gas, and other effects may change this.'),
        ('Replayed Battle card','Graveyard','Unless the replay effect or card says otherwise.'),
        ('Expired Condition or removed Asset','Discard pile','Unless an effect says otherwise.'),
    ]
    for a,b,c in rows:
        row=dest_table.add_row(); set_row_cant_split(row)
        row.cells[0].text=a; row.cells[1].text=b; row.cells[2].text=c
        for cell in row.cells: set_cell_margins(cell)
    for x in [
        'When the draw deck is empty, shuffle the discard pile to form a new deck. Graveyard cards are excluded.',
        'If the deck and discard pile together cannot provide the full number of cards, draw as many as possible.',
        "If multiple unresolved effects would send the same card to different destinations and neither explicitly overrides the other, place it in its owner's discard pile.",
        "A card's destination is resolved only once.",
    ]: add_bullet(doc,x)
    doc.add_heading('Assets', level=2)
    for x in [
        'An Asset is a persistent card in your Assets area. You may control no more than three.',
        'To play an Asset into a full area, you may discard one Asset you control to make room.',
        'If an effect would place a card into your full Assets area, you may discard an Asset to make room; otherwise the incoming card goes to its normal destination.',
        'A face-down Asset remains controlled and occupies a slot, but its printed effects are inactive.',
        'Face-down Assets can still be targeted or discarded. Turning one face up is not the same as placing a new Asset.',
        'Removed Assets normally go to the discard pile unless an effect says otherwise.',
    ]: add_bullet(doc,x)
    doc.add_heading('Conditions', level=2)
    for x in [
        'A Condition is a face-up card placed on a player until its printed trigger, duration, or removal condition ends.',
        'A Condition affects the player whose Conditions area contains it, but remains owned and controlled by the player who played it.',
        'Conditions do not occupy Asset slots.',
        'A player may have only one Condition with a given name. The controller of an incoming copy may replace the existing copy; otherwise the new card goes to its normal destination.',
        'Effects referring only to Assets do not affect Conditions.',
        'When a Condition expires, place it in its owner\'s discard pile unless its text says otherwise.',
    ]: add_bullet(doc,x)
    doc.add_heading('Overlays', level=2)
    for x in [
        'An Overlay is placed over a revealed Territory. Overlays cannot be placed on Heartlands.',
        'Only the top exposed Overlay is active. It supersedes the printed effect of the card immediately beneath it.',
        'Lower Overlays are dormant; their effects and expiration timers pause until they are exposed again.',
        'Changing control of a Territory does not remove or rearrange its Overlays.',
        'A removed Overlay normally goes to its owner\'s discard pile.',
        'A face-down Siege Weaponry card is an active Ruins Overlay when exposed. It continues to suppress the underlying Territory until specifically repaired.',
    ]: add_bullet(doc,x)

    # 7 special
    doc.add_heading('7. Special rules and winning', level=1)
    doc.add_heading('Mimicking and replaying cards', level=2)
    for x in [
        "When a card uses another card's Battle effect without moving that card, the other effect may be used only if its timing is still open and it can resolve independently.",
        'Cancellation, special-reveal, card-replacement, battle-ending, follow-up-battle, source-dependent, and card-lifecycle effects cannot be copied unless a card explicitly says otherwise.',
        'Contraband physically replays the chosen card. The replayed card follows its own text and lifecycle and counts as neither hand-origin nor battle-drawn unless an effect says otherwise.',
    ]: add_bullet(doc,x)
    doc.add_heading('Effect precedence', level=2)
    for x in [
        'Explicit override text takes precedence over ordinary effects.',
        'A cannot instruction overrides may and must unless another effect explicitly overrides the prohibition.',
        'Specific card text takes precedence over a general rule when it does not conflict with a prohibition.',
        'If two effects directly conflict and neither overrides the other, the prohibition takes precedence.',
    ]: add_bullet(doc,x)
    doc.add_heading('Manifest Destiny', level=2)
    for x in [
        'Manifest Destiny becomes a permanent blank Territory, not an Overlay.',
        'For its Action, insert it between your Heartland and your nearest Territory.',
        'For its Battle effect, insert it between the contested Territory and the space from which you attacked.',
        'Shift Territory cards as needed; tokens remain on their current spaces, and the insertion is not movement.',
        'The new Territory begins under your control, has no printed effect, counts as revealed for targeting, and can be occupied and captured normally.',
        'Heartlands do not move. The Action increases the opponent\'s victory distance by one.',
        'Manifest Destiny remains part of the board even if control changes.',
    ]: add_bullet(doc,x)
    doc.add_heading('Winning the game', level=2)
    for x in [
        'Advance through the Gauntlet and enter the opponent\'s Heartland.',
        'If the opponent\'s token is not in their Heartland, you win immediately upon entering.',
        'If the opponent occupies their Heartland, entering begins a battle. You must win that battle to win the game.',
        'When defending in your own Heartland, you make a Last Stand: you have Homeland Advantage and gain +1 to your battle total.',
        'If you lose a battle in the opponent\'s Heartland, retreat one tile toward your own Heartland.',
        'Heartlands are not Territories, cannot be controlled, and cannot receive Territory effects or Overlays unless a card explicitly says otherwise.',
    ]: add_bullet(doc,x)

    # 8 territories
    doc.add_page_break(); doc.add_heading('8. Territory reference', level=1)
    doc.add_paragraph('Choose exactly three different Territories. You may choose no more than one Arena. All Territory effects are active only while revealed unless a rule says otherwise.')
    ttable = doc.add_table(rows=0, cols=2); ttable.alignment=WD_TABLE_ALIGNMENT.CENTER; ttable.style='Table Grid'
    for i in range(0,len(TERRITORIES),2):
        row=ttable.add_row(); set_row_cant_split(row)
        for j in range(2):
            cell=row.cells[j]
            set_cell_margins(cell,95,110,95,110)
            if i+j < len(TERRITORIES):
                t=TERRITORIES[i+j]
                p=cell.paragraphs[0]; r=p.add_run(t['name']); r.bold=True; r.font.size=Pt(10)
                if t['name'].startswith('Arena:'): set_cell_shading(cell,'EFEFEF')
                p2=cell.add_paragraph(t['text']); p2.paragraph_format.space_after=Pt(0)
                for rr in p2.runs: rr.font.size=Pt(8.5)
            else:
                cell.text=''
    
    # 9 cards
    doc.add_page_break(); doc.add_heading('9. Canonical card reference', level=1)
    doc.add_paragraph('Every main-deck card has an Action effect and a Battle effect. Costs are provisional for playtesting. The source-specific battle-card lifecycle in Section 6 applies unless the card explicitly changes a destination.')
    ctable = doc.add_table(rows=0, cols=2); ctable.alignment=WD_TABLE_ALIGNMENT.CENTER; ctable.style='Table Grid'
    for i in range(0,len(CARDS),2):
        row=ctable.add_row(); set_row_cant_split(row)
        for j in range(2):
            cell=row.cells[j]
            if i+j < len(CARDS): add_card_reference_cell(cell,CARDS[i+j])
            else: cell.text=''

    # 10 decks
    doc.add_page_break(); doc.add_heading('10. Recommended playtest decks', level=1)
    doc.add_paragraph('These four 30-card decks provide distinct strategies for structured physical testing. Territory order is listed from the owner\'s Heartland toward the center. Do not alter a deck during the first round-robin unless recording it as a custom deck.')
    for name, deck in DECKS.items():
        doc.add_heading(name, level=2)
        p=doc.add_paragraph(); p.add_run(deck['subtitle']).italic=True
        p=doc.add_paragraph()
        p.add_run(f'Cards: 30 | Points: {deck_cost(deck["cards"])} | Territories: {", ".join(deck["territories"])}').bold=True
        doc.add_paragraph(deck['status'])
        counts=Counter(deck['cards'])
        table=doc.add_table(rows=1,cols=3); table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
        set_repeat_table_header(table.rows[0])
        for i,h in enumerate(['QTY','CARD','POINTS']): table.cell(0,i).text=h; set_cell_shading(table.cell(0,i),'D9D9D9')
        for card_name in sorted(counts):
            row=table.add_row(); set_row_cant_split(row)
            row.cells[0].text=str(counts[card_name]); row.cells[1].text=card_name; row.cells[2].text=str(counts[card_name]*CARD_BY_NAME[card_name]['cost'])
            for c in row.cells: set_cell_margins(c,45,75,45,75)

    # 11 playtest
    doc.add_page_break(); doc.add_heading('11. Playtest procedure and record sheet', level=1)
    doc.add_heading('Recommended first test cycle', level=2)
    for x in [
        'Play every pairing of the four recommended decks at least twice, switching which player uses each deck.',
        'Do not make balance changes between games in the same test cycle unless a rule is impossible to resolve.',
        'Record actual elapsed time and individual player-turn count. A round contains one turn by each player.',
        'Record every rules ambiguity verbatim before discussing a solution.',
        'After each game, each player independently rates enjoyment, agency, pacing, and desire to replay from 1 to 5.',
        'Flag cards and Territories that were never attractive to use, as well as effects that appeared mandatory whenever available.',
    ]: add_bullet(doc,x)
    doc.add_heading('Game record', level=2)
    fields = [
        ('Date / location',''),('Players',''),('Version',VERSION),('Deck A / player',''),('Deck B / player',''),('First player',''),('Winner',''),('Ending','Heartland entered / Heartland battle / concession / other'),('Elapsed time',''),('Player-turns',''),('Battles',''),('Captures',''),('Approx. reshuffles A / B',''),('Final Graveyard A / B',''),('Rules questions',''),('Most valuable card / Territory',''),('Least valuable card / Territory',''),('Was the outcome decided too early?',''),('Did either player feel unable to act?',''),('Enjoyment A / B (1-5)',''),('Agency A / B (1-5)',''),('Pacing A / B (1-5)',''),('Would replay? A / B',''),('Notes',''),
    ]
    table=doc.add_table(rows=0,cols=2); table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for label,default in fields:
        row=table.add_row(); set_row_cant_split(row)
        row.cells[0].text=label
        row.cells[0].paragraphs[0].runs[0].bold=True
        row.cells[1].text=default + ('\n\n' if label in ('Rules questions','Notes') else '')
        for c in row.cells: set_cell_margins(c,75,90,75,90)

    # Appendix quick reference
    doc.add_page_break(); doc.add_heading('Appendix A. Quick reference', level=1)
    quick = [
        ('TURN','Capture check - Draw 1 - one Action before or after Movement - Move - Cleanup to 3 cards.'),
        ('MOVEMENT','Advance 1, hold, or voluntarily withdraw 1. Entering an occupied space begins battle and ends unused movement.'),
        ('BATTLE CARDS','Up to 1 from hand and up to 1 from an initial draw of 3, unless modified.'),
        ('BATTLE','Reveal - cancellation - other effects - roll - rerolls - Revolution - modifiers - result.'),
        ('TIES','A defender on a controlled Territory wins. Otherwise reroll.'),
        ('DESTINATIONS','Hand-origin played card -> Graveyard. Battle-drawn played card -> discard. Replayed card -> Graveyard. Unused battle draw -> discard.'),
        ('CAPTURE','Win as attacker -> occupy. If still there at the start of your next turn -> capture, unless delayed.'),
        ('HAND','Draw 1 each turn; limit 3 at cleanup.'),
        ('ASSETS','Maximum 3. Face down = controlled but inactive.'),
        ('CONDITIONS','One of each name per player. Different names coexist.'),
        ('INCOMPLETE DRAW','Shuffle the discard pile when the deck empties, then draw as many cards as possible.'),
    ]
    for title,text in quick: add_rule_box(doc,title,text)

    # Appendix changes
    doc.add_heading('Appendix B. v0.5.4 change summary', level=1)
    doc.add_heading('Changes from v0.5.3', level=2)
    add_bullet(doc, 'Last Stand added: when defending in your own Heartland, you have Homeland Advantage and gain +1 to your battle total.')
    add_bullet(doc, 'Heartlands are explicitly not Territories, cannot be controlled, and cannot receive Territory effects or Overlays unless a card explicitly says otherwise.')
    add_bullet(doc, 'Heartland cards and player aids now include the Last Stand reminder.')
    doc.add_heading('Rules preserved from v0.5.3', level=2)
    add_bullet(doc, 'Battle-card destinations remain source-specific: hand-origin cards go to the Graveyard; battle-drawn cards recycle to the discard pile.')
    add_bullet(doc, 'The incomplete-draw rule instructs players to draw as many cards as possible.')
    add_bullet(doc, 'Attrition continues to convert an opponent\'s recyclable battle-draw cards into permanent losses after that opponent loses.')
    add_bullet(doc, 'Heartland remains the name of each player\'s endpoint and victory space.')
    doc.add_heading('Balance watch list', level=2)
    for x in ['Shock and Awe','Manifest Destiny','Siege Weaponry','Treason','Arcane Knowledge','Blockade','Protracted Siege','Attrition mirrors','High Ground','Supply Depot','Command Tent','Field Hospital']:
        add_bullet(doc,x)

    path = ROOT / 'Gauntlet_v0.5.4_Complete_Playtest_Guide.docx'
    doc.save(path)
    return path

# ---------------------------------------------------------------------------
# PDF card layout
# ---------------------------------------------------------------------------

PAGE_W, PAGE_H = letter
CARD_W, CARD_H = 2.5*inch, 3.5*inch
LEFT = 0.5*inch
BOTTOM = 0.25*inch

DECK_FILLS = {
    'Vanguard Coalition': colors.HexColor('#E8EDF2'),
    'Bastion Coalition': colors.HexColor('#E8E8E8'),
    'Reclamation Front': colors.HexColor('#EFEAE2'),
    'Shadow Directorate': colors.HexColor('#E9E5ED'),
}


def safe_html(s: str) -> str:
    return (s.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;'))


def fit_paragraph(c: canvas.Canvas, html: str, x: float, y: float, w: float, h: float, font_start=7.4, font_min=5.3, leading_ratio=1.16, italic=False) -> float:
    fs = font_start
    while fs >= font_min:
        style = ParagraphStyle(
            name='card', fontName='Helvetica-Oblique' if italic else 'Helvetica',
            fontSize=fs, leading=fs*leading_ratio, textColor=colors.black,
            alignment=TA_LEFT, spaceAfter=0, spaceBefore=0,
        )
        p = Paragraph(html, style)
        _, ph = p.wrap(w, h)
        if ph <= h:
            p.drawOn(c, x, y+h-ph)
            return fs
        fs -= 0.2
    style = ParagraphStyle(name='cardmin', fontName='Helvetica', fontSize=font_min, leading=font_min*1.08, alignment=TA_LEFT)
    p = Paragraph(html, style)
    kif = KeepInFrame(w, h, [p], mode='shrink')
    kif.wrapOn(c,w,h); kif.drawOn(c,x,y)
    return font_min


def draw_crop_marks(c: canvas.Canvas, x: float, y: float, w: float, h: float) -> None:
    c.setStrokeColor(colors.black); c.setLineWidth(0.35)
    m=0.10*inch
    # outside corners where space exists
    c.line(x-m,y,x,y); c.line(x,y-m,x,y)
    c.line(x+w,y,x+w+m,y); c.line(x+w,y-m,x+w,y)
    c.line(x-m,y+h,x,y+h); c.line(x,y+h,x,y+h+m)
    c.line(x+w,y+h,x+w+m,y+h); c.line(x+w,y+h,x+w,y+h+m)


def draw_main_card(c: canvas.Canvas, x: float, y: float, card: Dict[str,Any], deck_name: str|None=None, serial: str='') -> None:
    c.saveState()
    c.setFillColor(colors.white); c.setStrokeColor(colors.black); c.setLineWidth(0.7)
    c.rect(x,y,CARD_W,CARD_H,fill=1,stroke=1)
    header_h=0.43*inch
    fill = DECK_FILLS.get(deck_name, colors.HexColor('#ECECEC'))
    c.setFillColor(fill); c.rect(x,y+CARD_H-header_h,CARD_W,header_h,fill=1,stroke=0)
    c.setFillColor(colors.black)
    # Title dynamic size
    title=card['name']; title_fs=10.7
    maxw=CARD_W-0.62*inch
    while stringWidth(title,'Helvetica-Bold',title_fs)>maxw and title_fs>7.6: title_fs-=0.3
    c.setFont('Helvetica-Bold',title_fs)
    c.drawString(x+0.12*inch,y+CARD_H-0.28*inch,title)
    # cost badge
    bx=x+CARD_W-0.37*inch; by=y+CARD_H-0.22*inch
    c.setFillColor(colors.white); c.setStrokeColor(colors.black); c.circle(bx,by,0.15*inch,fill=1,stroke=1)
    c.setFillColor(colors.black); c.setFont('Helvetica-Bold',10); c.drawCentredString(bx,by-3.5,str(card['cost']))
    if card.get('unique'):
        c.setFont('Helvetica-Bold',5.5); c.drawRightString(x+CARD_W-0.10*inch,y+CARD_H-header_h-0.10*inch,'UNIQUE')
    # body divider
    body_x=x+0.12*inch; body_w=CARD_W-0.24*inch
    footer_h=0.22*inch
    body_bottom=y+footer_h+0.06*inch
    body_top=y+CARD_H-header_h-0.08*inch
    total_h=body_top-body_bottom
    action_len=len(card['action']); battle_len=len(card['battle']); rem_len=len(card.get('reminder',''))
    rem_h = 0.34*inch if rem_len else 0
    usable=total_h-rem_h-0.10*inch
    ratio=max(0.39,min(0.61,(action_len+40)/(action_len+battle_len+80)))
    action_h=usable*ratio
    battle_h=usable-action_h
    c.setStrokeColor(colors.HexColor('#888888')); c.setLineWidth(0.35)
    split_y=body_bottom+battle_h+rem_h+0.05*inch
    c.line(body_x,split_y,body_x+body_w,split_y)
    # labels and text
    c.setFont('Helvetica-Bold',6.8); c.setFillColor(colors.black)
    c.drawString(body_x,split_y+action_h-0.12*inch,'ACTION')
    fit_paragraph(c,safe_html(card['action']),body_x,split_y+0.03*inch,body_w,action_h-0.18*inch,7.1,5.3)
    battle_top=split_y-0.02*inch
    c.setFont('Helvetica-Bold',6.8); c.drawString(body_x,battle_top-0.10*inch,'BATTLE')
    fit_paragraph(c,safe_html(card['battle']),body_x,body_bottom+rem_h+0.02*inch,body_w,battle_h-0.16*inch,7.1,5.3)
    if rem_len:
        c.setStrokeColor(colors.HexColor('#B0B0B0')); c.line(body_x,body_bottom+rem_h,body_x+body_w,body_bottom+rem_h)
        fit_paragraph(c,'<b><i>REMINDER:</i></b> <i>'+safe_html(card['reminder'])+'</i>',body_x,body_bottom,body_w,rem_h-0.02*inch,6.0,4.9,1.08,False)
    # footer
    c.setFillColor(colors.HexColor('#F2F2F2')); c.rect(x,y,CARD_W,footer_h,fill=1,stroke=0)
    c.setFillColor(colors.black); c.setFont('Helvetica',5.3)
    footer=(deck_name.upper() if deck_name else 'MASTER POOL')
    c.drawString(x+0.08*inch,y+0.075*inch,footer)
    c.drawRightString(x+CARD_W-0.08*inch,y+0.075*inch,f'{VERSION}{("  "+serial) if serial else ""}')
    draw_crop_marks(c,x,y,CARD_W,CARD_H)
    c.restoreState()


def draw_territory_card(c: canvas.Canvas, x: float, y: float, territory: Dict[str,str], deck_name: str|None=None, serial: str='') -> None:
    c.saveState(); c.setFillColor(colors.white); c.setStrokeColor(colors.black); c.setLineWidth(0.9)
    c.rect(x,y,CARD_W,CARD_H,fill=1,stroke=1)
    is_arena=territory['name'].startswith('Arena:')
    fill=colors.HexColor('#D9D9D9') if not is_arena else colors.HexColor('#C8C8C8')
    c.setFillColor(fill); c.rect(x,y+CARD_H-0.62*inch,CARD_W,0.62*inch,fill=1,stroke=0)
    c.setFillColor(colors.black); c.setFont('Helvetica-Bold',6.5)
    c.drawString(x+0.12*inch,y+CARD_H-0.16*inch,'ARENA TERRITORY' if is_arena else 'TERRITORY')
    title=territory['name']; fs=12
    maxw=CARD_W-0.24*inch
    while stringWidth(title,'Helvetica-Bold',fs)>maxw and fs>7.5: fs-=0.3
    c.setFont('Helvetica-Bold',fs); c.drawCentredString(x+CARD_W/2,y+CARD_H-0.44*inch,title)
    fit_paragraph(c,safe_html(territory['text']),x+0.15*inch,y+0.42*inch,CARD_W-0.30*inch,CARD_H-1.20*inch,9.0,6.3,1.18)
    c.setFillColor(colors.HexColor('#F2F2F2')); c.rect(x,y,CARD_W,0.28*inch,fill=1,stroke=0)
    c.setFillColor(colors.black); c.setFont('Helvetica',5.3)
    c.drawString(x+0.08*inch,y+0.09*inch,(deck_name.upper() if deck_name else 'MASTER POOL'))
    c.drawRightString(x+CARD_W-0.08*inch,y+0.09*inch,f'{VERSION}{("  "+serial) if serial else ""}')
    draw_crop_marks(c,x,y,CARD_W,CARD_H); c.restoreState()


def draw_heartland_card(c:canvas.Canvas,x:float,y:float,label:str='HEARTLAND',deck_name:str|None=None) -> None:
    c.saveState(); c.setFillColor(colors.white); c.setStrokeColor(colors.black); c.setLineWidth(1.0)
    c.rect(x,y,CARD_W,CARD_H,fill=1,stroke=1)
    c.setFillColor(colors.HexColor('#D5D5D5')); c.rect(x,y+CARD_H-0.75*inch,CARD_W,0.75*inch,fill=1,stroke=0)
    c.setFillColor(colors.black); c.setFont('Helvetica-Bold',16); c.drawCentredString(x+CARD_W/2,y+CARD_H-0.47*inch,label)
    fit_paragraph(c,'<b>Your force begins here.</b><br/><br/>An opponent wins by entering this Heartland while it is empty, or by entering and defeating your token here. When you defend here, you have Homeland Advantage and +1.',x+0.25*inch,y+0.65*inch,CARD_W-0.5*inch,CARD_H-1.65*inch,10,7.5,1.25)
    c.setFont('Helvetica-Bold',7); c.drawCentredString(x+CARD_W/2,y+0.35*inch,'HEARTLAND IS NOT A TERRITORY')
    c.setFont('Helvetica',5.3); c.drawCentredString(x+CARD_W/2,y+0.12*inch,f'{VERSION}  {deck_name or "PLAYTEST"}')
    draw_crop_marks(c,x,y,CARD_W,CARD_H); c.restoreState()


def draw_aid_card(c:canvas.Canvas,x:float,y:float,title:str,lines:List[str],deck_name:str|None=None) -> None:
    c.saveState(); c.setFillColor(colors.white); c.setStrokeColor(colors.black); c.setLineWidth(0.8)
    c.rect(x,y,CARD_W,CARD_H,fill=1,stroke=1)
    c.setFillColor(colors.HexColor('#D9D9D9')); c.rect(x,y+CARD_H-0.50*inch,CARD_W,0.50*inch,fill=1,stroke=0)
    c.setFillColor(colors.black); c.setFont('Helvetica-Bold',11); c.drawCentredString(x+CARD_W/2,y+CARD_H-0.31*inch,title)
    html='<br/>'.join(f'<b>{i+1}.</b> {safe_html(line)}' for i,line in enumerate(lines))
    fit_paragraph(c,html,x+0.15*inch,y+0.35*inch,CARD_W-0.30*inch,CARD_H-1.0*inch,8.2,6.2,1.25)
    c.setFont('Helvetica',5.3); c.drawCentredString(x+CARD_W/2,y+0.12*inch,f'{VERSION}  {deck_name or "PLAYER AID"}')
    draw_crop_marks(c,x,y,CARD_W,CARD_H); c.restoreState()


def page_positions() -> List[Tuple[float,float]]:
    # order top-left to bottom-right
    pos=[]
    for row in range(3):
        y=PAGE_H-BOTTOM-CARD_H*(row+1)
        for col in range(3):
            x=LEFT+CARD_W*col
            pos.append((x,y))
    return pos


def draw_card_pages(path:Path, items:List[Dict[str,Any]], title_meta:str) -> None:
    c=canvas.Canvas(str(path),pagesize=letter)
    positions=page_positions()
    for pidx in range(math.ceil(len(items)/9)):
        page_items=items[pidx*9:(pidx+1)*9]
        for idx,item in enumerate(page_items):
            x,y=positions[idx]
            kind=item['kind']
            if kind=='main': draw_main_card(c,x,y,item['card'],item.get('deck'),item.get('serial',''))
            elif kind=='territory': draw_territory_card(c,x,y,item['territory'],item.get('deck'),item.get('serial',''))
            elif kind=='heartland': draw_heartland_card(c,x,y,item.get('label','HEARTLAND'),item.get('deck'))
            elif kind=='aid': draw_aid_card(c,x,y,item['title'],item['lines'],item.get('deck'))
        # Print metadata outside cut area, tiny at bottom center
        c.setFont('Helvetica',5); c.setFillColor(colors.HexColor('#666666'))
        c.drawCentredString(PAGE_W/2,0.07*inch,f'{title_meta} - sheet {pidx+1} of {math.ceil(len(items)/9)}')
        c.showPage()
    c.save()


def build_four_deck_pdf() -> Path:
    items=[]
    for deck_name,deck in DECKS.items():
        seq=Counter()
        for name in deck['cards']:
            seq[name]+=1
            items.append({'kind':'main','card':CARD_BY_NAME[name],'deck':deck_name,'serial':f'{seq[name]}'})
        for t in deck['territories']:
            items.append({'kind':'territory','territory':TERRITORY_BY_NAME[t],'deck':deck_name})
        items.append({'kind':'heartland','deck':deck_name})
        items.append({'kind':'aid','deck':deck_name,'title':'TURN SEQUENCE','lines':[
            'Capture check.','Draw one card.','Optional Action before movement.','Advance, hold, or withdraw.','Optional Action after movement if unused.','Cleanup to three cards.'
        ]})
        items.append({'kind':'aid','deck':deck_name,'title':'BATTLE & DESTINATIONS','lines':[
            'Optional one card from hand.','Draw three; select up to one.','Reveal; cancel; resolve effects.','Roll and determine winner.','Loser retreats; attacker may occupy.','Hand to GY; battle-draw to discard. Heartland: ties/+1.'
        ]})
    # 36 cards/deck = 144 cards = 16 sheets
    assert len(items)==144, len(items)
    path=ROOT/'Gauntlet_v0.5.4_Four_Deck_Playtest_Set.pdf'
    draw_card_pages(path,items,'Gauntlet v0.5.4 Four-Deck Playtest Set')
    return path


def build_master_pool_pdf() -> Path:
    items=[]
    for card in CARDS: items.append({'kind':'main','card':card})
    for t in TERRITORIES: items.append({'kind':'territory','territory':t})
    items.append({'kind':'heartland','label':'HEARTLAND A'})
    items.append({'kind':'heartland','label':'HEARTLAND B'})
    # 81 exactly (54+25+2)
    assert len(items)==81
    path=ROOT/'Gauntlet_v0.5.4_Master_Card_Pool.pdf'
    draw_card_pages(path,items,'Gauntlet v0.5.4 Master Card Pool')
    return path


def build_playtest_forms_pdf() -> Path:
    path=ROOT/'Gauntlet_v0.5.4_Playtest_Forms.pdf'
    c=canvas.Canvas(str(path),pagesize=letter)
    for page in range(4):
        margin=0.55*inch
        c.setFont('Helvetica-Bold',18); c.drawString(margin,PAGE_H-margin,'GAUNTLET PLAYTEST RECORD')
        c.setFont('Helvetica-Bold',10); c.drawRightString(PAGE_W-margin,PAGE_H-margin+2,VERSION)
        y=PAGE_H-margin-0.38*inch
        fields=[
            ('Date / location',2.0),('Players',2.0),('Deck A / player',2.4),('Deck B / player',2.4),
            ('First player',1.8),('Winner',1.8),('Ending',2.8),('Elapsed time',1.8),('Player-turns',1.8),
            ('Battles',1.4),('Captures',1.4),('Reshuffles A / B',2.0),('Final Graveyard A / B',2.3),
        ]
        c.setFont('Helvetica',8)
        colx=[margin,PAGE_W/2+0.12*inch]
        idx=0
        for row in range(7):
            for col in range(2):
                if idx>=len(fields): break
                label,w=fields[idx]; x=colx[col]
                c.setFont('Helvetica-Bold',7.5); c.drawString(x,y,label.upper())
                c.line(x,y-0.08*inch,x+(PAGE_W/2-margin-0.2*inch),y-0.08*inch)
                idx+=1
            y-=0.42*inch
        # ratings
        y-=0.05*inch
        c.setFont('Helvetica-Bold',9); c.drawString(margin,y,'PLAYER RATINGS (1 = poor, 5 = excellent)')
        y-=0.25*inch
        c.setFont('Helvetica',8)
        for label in ['Enjoyment','Agency','Pacing','Desire to replay']:
            c.drawString(margin,y,label+' - Player A:  1  2  3  4  5        Player B:  1  2  3  4  5')
            y-=0.25*inch
        prompts=[
            'Rules questions or ambiguous interactions',
            'Most valuable card or Territory',
            'Least valuable card or Territory',
            'Did either player feel unable to act? Why?',
            'Was the winner apparent long before the game ended?',
            'Other observations',
        ]
        for prompt in prompts:
            c.setFont('Helvetica-Bold',8); c.drawString(margin,y,prompt.upper()); y-=0.13*inch
            lines=2 if prompt!='Other observations' else 4
            for _ in range(lines):
                c.line(margin,y,PAGE_W-margin,y); y-=0.22*inch
            y-=0.07*inch
        c.setFont('Helvetica',5.5); c.drawCentredString(PAGE_W/2,0.25*inch,f'Gauntlet {VERSION} - Record sheet {page+1} of 4')
        c.showPage()
    c.save(); return path


def create_readme() -> Path:
    text=f'''# Gauntlet {VERSION} Print-and-Play Package

Generated {DATE}.

## Files

- `Gauntlet_v0.5.4_Complete_Playtest_Guide.docx` - editable canonical guide.
- `Gauntlet_v0.5.4_Complete_Playtest_Guide.pdf` - print-ready guide.
- `Gauntlet_v0.5.4_Four_Deck_Playtest_Set.pdf` - four complete 30-card decks, twelve Territories, four Heartlands, and eight player aids.
- `Gauntlet_v0.5.4_Master_Card_Pool.pdf` - one copy of all 54 main-deck cards, all 25 Territories, and two Heartlands.
- `Gauntlet_v0.5.4_Playtest_Forms.pdf` - four game record sheets.
- `Gauntlet_v0.5.4_Canonical_Data.json` - editable structured source data.
- `Gauntlet_v0.5.4_Playtest_Deck_Lists.md` - deck manifests.

## Printing

Card sheets use exact 2.5 x 3.5 inch poker-card fronts in a 3 x 3 US Letter layout. Print at **Actual Size / 100%**, with scaling disabled. For a durable prototype, cut the fronts and sleeve them in opaque standard poker sleeves with ordinary playing cards or bulk trading cards behind them. This avoids duplex alignment problems.

The Four-Deck Playtest Set is organized deck-by-deck and labels every card at the bottom for sorting. The Shadow Directorate is the v0.5.4 Attrition trial list; Attrition replaces Court Martial at equal cost.

## Rules changes incorporated

- Heartland terminology.
- Hand-origin battle cards go to the Graveyard.
- Battle-drawn battle cards go to the discard pile.
- Replayed Battle cards normally go to the Graveyard.
- Incomplete draws draw as many cards as possible.
- Attrition uses the locked three-card Battle effect.
- Last Stand: Heartland defenders have Homeland Advantage and gain +1.
'''
    path=ROOT/'README.md'; path.write_text(text,encoding='utf-8'); return path


def make_zip(paths: List[Path]) -> Path:
    zpath=ROOT/'Gauntlet_v0.5.4_Print_and_Play_Package.zip'
    with zipfile.ZipFile(zpath,'w',zipfile.ZIP_DEFLATED) as z:
        for p in paths:
            if p.exists(): z.write(p,arcname=p.name)
    return zpath


def main():
    validate()
    outputs=[]
    outputs.append(export_json())
    outputs.append(export_deck_lists_md())
    outputs.append(build_docx())
    outputs.append(build_four_deck_pdf())
    outputs.append(build_master_pool_pdf())
    outputs.append(build_playtest_forms_pdf())
    outputs.append(create_readme())
    print('\n'.join(str(p) for p in outputs))
    # Zip made after DOCX->PDF conversion externally.

if __name__=='__main__':
    main()
