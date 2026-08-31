#!/usr/bin/env python3
"""Build the v0.6.1 Rulebook and Reference Guide as DOCX and PDF.

Requirements available in CI:
- pandoc
- LibreOffice
- python-docx

The Markdown governing sources remain authoritative. Generated DOCX/PDF files are
release artifacts and must not be edited independently.
"""

from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
RELEASE = ROOT / "releases" / "v0.6.1"
RULEBOOK_MD = RELEASE / "Gauntlet_v0.6.1_Rulebook.md"
REFERENCE_MD = RELEASE / "Gauntlet_v0.6.1_Reference_Guide.md"

ACCENT = RGBColor(111, 53, 24)
INK = RGBColor(36, 27, 23)
MUTED = RGBColor(99, 86, 76)


def require(command: str) -> str:
    path = shutil.which(command)
    if not path:
        raise RuntimeError(f"Required command not found: {command}")
    return path


def set_run_font(run, name: str, size: float, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def ensure_style(document: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, style_type)


def build_reference_doc(path: Path, short_title: str) -> None:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)

    normal = document.styles["Normal"]
    normal.font.name = "Liberation Sans"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Liberation Sans")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = 1.05

    title = document.styles["Title"]
    title.font.name = "Liberation Serif"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Liberation Serif")
    title.font.size = Pt(27)
    title.font.bold = True
    title.font.color.rgb = ACCENT
    title.paragraph_format.space_after = Pt(10)

    subtitle = ensure_style(document, "Subtitle")
    subtitle.font.name = "Liberation Sans"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Liberation Sans")
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(10)

    for index, (size, before, after) in enumerate(
        [(18, 15, 6), (14, 11, 4), (11.5, 8, 3)], start=1
    ):
        style = document.styles[f"Heading {index}"]
        style.font.name = "Liberation Serif" if index < 3 else "Liberation Sans"
        style._element.rPr.rFonts.set(
            qn("w:eastAsia"), "Liberation Serif" if index < 3 else "Liberation Sans"
        )
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = document.styles[style_name]
        style.font.name = "Liberation Sans"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Liberation Sans")
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(2.5)

    quote = document.styles["Quote"]
    quote.font.name = "Liberation Serif"
    quote._element.rPr.rFonts.set(qn("w:eastAsia"), "Liberation Serif")
    quote.font.size = Pt(10.2)
    quote.font.italic = False
    quote.font.color.rgb = INK
    quote.paragraph_format.left_indent = Inches(0.25)
    quote.paragraph_format.right_indent = Inches(0.15)
    quote.paragraph_format.space_before = Pt(4)
    quote.paragraph_format.space_after = Pt(6)

    table = ensure_style(document, "Table", WD_STYLE_TYPE.TABLE)
    table.font.name = "Liberation Sans"
    table._element.rPr.rFonts.set(qn("w:eastAsia"), "Liberation Sans")
    table.font.size = Pt(8.8)

    compact = ensure_style(document, "Compact")
    compact.base_style = normal
    compact.paragraph_format.space_after = Pt(1.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(f"GAUNTLET  |  {short_title.upper()}  |  v0.6.1")
    set_run_font(run, "Liberation Sans", 7.5, bold=True)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Gauntlet v0.6.1  |  First Playtest Revision  |  Page ")
    set_run_font(run, "Liberation Sans", 7.5)
    run.font.color.rgb = MUTED
    add_page_number(footer)

    document.add_paragraph("Reference template - generated; do not edit.")
    document.save(path)


def run(command: list[str], cwd: Path | None = None) -> None:
    print("+", " ".join(command), flush=True)
    subprocess.run(command, cwd=cwd, check=True)


def convert_markdown(
    markdown: Path,
    output_docx: Path,
    output_pdf: Path,
    title: str,
    reference_doc: Path,
    include_toc: bool,
) -> None:
    pandoc = require("pandoc")
    libreoffice = require("libreoffice")

    command = [
        pandoc,
        str(markdown),
        "--from=gfm",
        "--to=docx",
        f"--reference-doc={reference_doc}",
        f"--metadata=title:{title}",
        "--metadata=subtitle:Version 0.6.1 - First Playtest Revision",
        "--standalone",
        "--wrap=none",
        "-o",
        str(output_docx),
    ]
    if include_toc:
        command.extend(["--toc", "--toc-depth=2"])
    run(command, cwd=ROOT)

    if include_toc:
        run(
            [
                sys.executable,
                str(ROOT / "scripts" / "finalize_v061_docx.py"),
                str(output_docx),
            ],
            cwd=ROOT,
        )

    with tempfile.TemporaryDirectory(prefix="gauntlet-lo-") as profile:
        run(
            [
                libreoffice,
                "--headless",
                f"-env:UserInstallation=file://{profile}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_pdf.parent),
                str(output_docx),
            ],
            cwd=ROOT,
        )

    generated = output_pdf.parent / f"{output_docx.stem}.pdf"
    if generated != output_pdf:
        generated.replace(output_pdf)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--check-tools",
        action="store_true",
        help="Check required external programs and exit.",
    )
    args = parser.parse_args()

    require("pandoc")
    require("libreoffice")
    if args.check_tools:
        print("Document build tools are available.")
        return 0

    if not RULEBOOK_MD.is_file() or not REFERENCE_MD.is_file():
        raise RuntimeError("Missing v0.6.1 Markdown sources")

    RELEASE.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="gauntlet-docx-") as temp:
        temp_dir = Path(temp)
        rulebook_template = temp_dir / "rulebook-reference.docx"
        reference_template = temp_dir / "guide-reference.docx"
        build_reference_doc(rulebook_template, "Official Rulebook")
        build_reference_doc(reference_template, "Reference Guide")

        convert_markdown(
            RULEBOOK_MD,
            RELEASE / "Gauntlet_v0.6.1_Rulebook.docx",
            RELEASE / "Gauntlet_v0.6.1_Rulebook.pdf",
            "GAUNTLET - Official Rulebook",
            rulebook_template,
            include_toc=True,
        )
        convert_markdown(
            REFERENCE_MD,
            RELEASE / "Gauntlet_v0.6.1_Reference_Guide.docx",
            RELEASE / "Gauntlet_v0.6.1_Reference_Guide.pdf",
            "GAUNTLET - Reference Guide",
            reference_template,
            include_toc=False,
        )

    for output in [
        RELEASE / "Gauntlet_v0.6.1_Rulebook.docx",
        RELEASE / "Gauntlet_v0.6.1_Rulebook.pdf",
        RELEASE / "Gauntlet_v0.6.1_Reference_Guide.docx",
        RELEASE / "Gauntlet_v0.6.1_Reference_Guide.pdf",
    ]:
        if not output.is_file() or output.stat().st_size < 10_000:
            raise RuntimeError(f"Output missing or unexpectedly small: {output}")
        print(f"Wrote {output.relative_to(ROOT)} ({output.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
