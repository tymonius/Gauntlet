from __future__ import annotations

from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]

REPLACEMENTS: dict[str, list[tuple[str, str]]] = {
    "README.md": [
        ("| Arcane | Rites, sacrifice, transformation, and Ritual victory |",
         "| Mystics | Rites, sacrifice, transformation, and Ritual victory |"),
    ],
    "docs/Gauntlet_v0.6_Open_Questions.md": [
        ("Complete Arcane, Financier, and Intelligence twelve-card packages.",
         "Complete the Mystics, Financiers, and Intelligence twelve-card packages."),
        ("Does Blasphemy pressure Arcane without creating a hard counter?",
         "Does Blasphemy pressure Mystic decks without creating a hard counter?"),
        ("## Arcane\n", "## Mystics\n"),
    ],
    "docs/Gauntlet_v0.6_Card_Metadata.md": [
        ("| 1 | Witchcraft | Arcane | TBD | Advanced |",
         "| 1 | Witchcraft | Mystics | TBD | Advanced |"),
        ("| 27 | Necromancy | Arcane | TBD | Advanced | Recursion and Arcane-system interactions |",
         "| 27 | Necromancy | Mystics | TBD | Advanced | Recursion and interactions with Mystics mechanics |"),
        ("Remaining card work concerns the Arcane, Financier, and Intelligence packages,",
         "Remaining card work concerns the Mystics, Financiers, and Intelligence packages,"),
    ],
    "docs/Gauntlet_v0.7_Parking_Lot.md": [
        ("| Diplomats | Arcane |", "| Diplomats | Mystics |"),
        ("- Arcane during Night: first Arcane card sent from hand to Graveyard each turn draws one,",
         "- Mystics during Night: first Arcane card sent from hand to Graveyard each turn draws one,"),
    ],
    "docs/Gauntlet_Design_Principles_and_Guardrails.md": [
        ("- **Arcane:** Rites, sacrifice, transformation, Graveyard interaction, and Ritual.",
         "- **Mystics:** Rites, sacrifice, transformation, Graveyard interaction, and Ritual."),
        ("Must have a functional plan against every faction, not only Arcane.",
         "Must have a functional plan against every faction, not only Mystics."),
        ("### Arcane\n", "### Mystics\n"),
        ("- Arcane may be advanced, but should not become a solitaire puzzle.",
         "- Mystics may be advanced, but should not become a solitaire puzzle."),
    ],
    "docs/Gauntlet_v0.6_Card_Review_Log.md": [
        ("**Allegiance:** Arcane\n", "**Allegiance:** Mystics\n"),
        ("**Allegiance:** Arcane  \n", "**Allegiance:** Mystics  \n"),
        ("interactions with Arcane Graveyard pressure and Inquisition Purification.",
         "interactions with Graveyard pressure from Mystics and Inquisition Purification."),
    ],
    "docs/territory-reviews/017-monastery.md": [
        ("**Watchlist:** Arcane-faction matchups and Arcane-trait assignment",
         "**Watchlist:** Matchups involving Mystics and final Arcane-trait assignment"),
        ("without inherently disabling Arcane faction rules or leader abilities.",
         "without inherently disabling Mystics rules or leader abilities."),
    ],
    "docs/territory-reviews/STATUS.md": [
        ("- Arcane-faction pressure from Monastery and final Arcane-trait assignment.",
         "- Pressure on Mystics from Monastery and final Arcane-trait assignment."),
    ],
    "docs/territory-reviews/014-old-battlefield.md": [
        ("**Watchlist:** Narrow utility, Arcane Graveyard setup, and deliberate deck thinning",
         "**Watchlist:** Narrow utility, Graveyard setup by Mystics, and deliberate deck thinning"),
    ],
    "docs/Gauntlet_v0.6_Project_Index.md": [
        ("assigned to Arcane, Financiers, and Intelligence.",
         "assigned to Mystics, Financiers, and Intelligence."),
        ("Complete Arcane, Financier, and Intelligence card packages.",
         "Complete the Mystics, Financiers, and Intelligence card packages."),
        ("- Arcane — Alchemist / Spirit Walker",
         "- Mystics — Alchemist / Spirit Walker"),
        ("- Arcane, Financiers, and Intelligence remain disabled placeholders until their packages are complete.",
         "- Mystics, Financiers, and Intelligence remain disabled placeholders until their packages are complete."),
        ("while continuing Arcane, Financier, or Intelligence design.",
         "while continuing design work on Mystics, Financiers, or Intelligence."),
    ],
    "docs/card-reviews/STATUS.md": [
        ("assigned to Arcane, Financiers, and Intelligence.",
         "assigned to Mystics, Financiers, and Intelligence."),
        ("- Arcane trait and allegiance are separate.",
         "- Arcane trait and faction allegiance are separate."),
        ("- Complete Arcane, Financier, and Intelligence packages.",
         "- Complete the Mystics, Financiers, and Intelligence packages."),
        ("Begin Arcane, Financier, or Intelligence design",
         "Begin design work on Mystics, Financiers, or Intelligence"),
    ],
    "docs/card-reviews/054-arcane-knowledge.md": [
        ("cost 5, Arcane allegiance, and Advanced complexity:",
         "cost 5, allegiance to Mystics, and Advanced complexity:"),
        ("not automatically locked behind Arcane allegiance.",
         "not automatically restricted to Mystics."),
        ("while the Arcane faction retains the highest density and strongest systemic exploitation of such cards.",
         "while the Mystics faction retains the highest density and strongest systemic exploitation of such cards."),
        ("outside a direct Arcane-faction matchup.",
         "outside a direct matchup against Mystics."),
        ("The Arcane faction can exploit it more efficiently",
         "Mystics can exploit it more efficiently"),
        ("without granting access to the Arcane faction's full engine.",
         "without granting access to the Mystics' full faction engine."),
        ("Test Arcane Knowledge in non-Arcane decks and in Arcane decks",
         "Test Arcane Knowledge in non-Mystic decks and in Mystic decks"),
    ],
    "docs/Gauntlet_v0.6_Working_Rules.md": [
        ("| **Arcane** | Rites, sacrifice, forbidden knowledge, transformation, and Ritual victory |",
         "| **Mystics** | Rites, sacrifice, forbidden knowledge, transformation, and Ritual victory |"),
        ("## 12. Arcane\n", "## 12. Mystics\n"),
        ("| Arcane | Pentagram or pentacle |", "| Mystics | Pentagram or pentacle |"),
    ],
    "docs/Gauntlet_v0.6_Territory_Pool.md": [
        ("**Watchlist:** Narrow utility, Arcane Graveyard setup, and deliberate deck thinning",
         "**Watchlist:** Narrow utility, Graveyard setup by Mystics, and deliberate deck thinning"),
        ("**Watchlist:** Arcane-faction matchups and Arcane-trait assignment",
         "**Watchlist:** Matchups involving Mystics and final Arcane-trait assignment"),
        ("Evaluate Monastery's pressure on the completed Arcane faction and confirm final Arcane-trait assignments.",
         "Evaluate Monastery's pressure on the completed Mystics faction and confirm final Arcane-trait assignments."),
    ],
    "docs/Gauntlet_v0.6_Leader_Design_Bible.md": [
        ("| Arcane | Bone, charcoal, dark red, green glass, brass, smoke, ritual marks |",
         "| Mystics | Bone, charcoal, dark red, green glass, brass, smoke, ritual marks |"),
        ("# Arcane Leaders\n", "# Mystic Leaders\n"),
        ("**Faction:** Arcane  \n", "**Faction:** Mystics  \n"),
        ("| Arcane | Alchemist | Spirit Walker | Experiment vs invocation |",
         "| Mystics | Alchemist | Spirit Walker | Experiment vs invocation |"),
        ("| Arcane | Lab apparatus / ritual circle |",
         "| Mystics | Lab apparatus / ritual circle |"),
    ],
    "docs/Gauntlet_Lore_Development_Notes.md": [
        ("- Arcane;\n", "- Mystics;\n"),
        ("with Arcane scholars while opposing them;",
         "with Mystic scholars while opposing them;"),
        ("How public and accepted is Arcane practice?",
         "How public and accepted are Mystic practices?"),
    ],
    "releases/v0.6/faction-guides/inquisition/Gauntlet_v0.6_Inquisition_Faction_Guide.md": [
        ("The Arcane trait is distinct from Arcane faction allegiance.",
         "The Arcane trait is separate from faction allegiance; not every Arcane card belongs to Mystics."),
    ],
    "releases/v0.6/Gauntlet_v0.6_Preliminary_Core_Rules.md": [
        ("Arcane, Financiers, and Intelligence remain in development.",
         "Mystics, Financiers, and Intelligence remain in development."),
    ],
    "deckbuilder-v0.6/app.js": [
        ("id: \"arcane\",\n    name: \"Arcane\",",
         "id: \"mystics\",\n    name: \"Mystics\","),
    ],
    "deckbuilder-v0.6/README.md": [
        ("Arcane, Financiers, and Intelligence appear as disabled development placeholders.",
         "Mystics, Financiers, and Intelligence appear as disabled development placeholders."),
        ("Add completed Arcane, Financier, and Intelligence packages",
         "Add the completed Mystics, Financiers, and Intelligence packages"),
    ],
    "deckbuilder-v0.6/supplemental-data.js": [
        ("The Arcane trait is distinct from Arcane faction allegiance.",
         "The Arcane trait is separate from faction allegiance; not every Arcane card belongs to Mystics."),
    ],
    "faction-sheets/inquisition.js": [
        ("The Arcane trait is distinct from Arcane faction allegiance.",
         "The Arcane trait is separate from faction allegiance; not every Arcane card belongs to Mystics."),
    ],
}


def replace_exact(path: Path, old: str, new: str) -> None:
    text = path.read_text(encoding="utf-8")
    count = text.count(old)
    if count == 0:
        raise RuntimeError(f"Missing expected text in {path.relative_to(ROOT)}: {old!r}")
    path.write_text(text.replace(old, new), encoding="utf-8")


changed: list[str] = []
for relative, pairs in REPLACEMENTS.items():
    path = ROOT / relative
    before = path.read_text(encoding="utf-8")
    for old, new in pairs:
        replace_exact(path, old, new)
    if path.read_text(encoding="utf-8") != before:
        changed.append(relative)

docx_path = ROOT / "releases/v0.6/faction-guides/inquisition/Gauntlet_v0.6_Inquisition_Faction_Guide.docx"
doc = Document(docx_path)
old_docx = "The Arcane trait is distinct from Arcane allegiance."
new_docx = "The Arcane trait is separate from faction allegiance; not every Arcane card belongs to Mystics."
docx_replacements = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if old_docx in paragraph.text:
                    for run in paragraph.runs:
                        if old_docx in run.text:
                            run.text = run.text.replace(old_docx, new_docx)
                            docx_replacements += 1
                            break
                    else:
                        paragraph.text = paragraph.text.replace(old_docx, new_docx)
                        docx_replacements += 1
if docx_replacements != 1:
    raise RuntimeError(f"Expected one DOCX replacement, found {docx_replacements}")
doc.save(docx_path)
changed.append(str(docx_path.relative_to(ROOT)))

forbidden = [
    "Arcane faction",
    "Arcane-faction",
    "Arcane allegiance",
    "**Allegiance:** Arcane",
    "| Arcane | Rites, sacrifice, transformation, and Ritual victory |",
    "| **Arcane** | Rites, sacrifice, forbidden knowledge, transformation, and Ritual victory |",
    "- Arcane — Alchemist / Spirit Walker",
    'id: "arcane"',
    'name: "Arcane"',
]
search_roots = [
    ROOT / "README.md",
    ROOT / "docs",
    ROOT / "releases/v0.6",
    ROOT / "deckbuilder-v0.6",
    ROOT / "faction-sheets",
]
violations: list[str] = []
for search_root in search_roots:
    paths = [search_root] if search_root.is_file() else search_root.rglob("*")
    for path in paths:
        if not path.is_file() or path.suffix.lower() in {".pdf", ".docx", ".png", ".jpg", ".jpeg", ".webp"}:
            continue
        if "Archive" in path.parts or "v0.5.7" in path.parts:
            continue
        text = path.read_text(encoding="utf-8", errors="ignore")
        for phrase in forbidden:
            if phrase in text:
                violations.append(f"{path.relative_to(ROOT)}: {phrase}")
if violations:
    raise RuntimeError("Remaining faction-name references:\n" + "\n".join(violations))

print(f"Updated {len(changed)} source files:")
for item in changed:
    print(f"- {item}")
