#!/usr/bin/env python3
"""Build the v0.6.1 Rulebook and Reference Guide documents.

The governing Markdown remains authoritative. The rulebook DOCX is an editable
half-letter edition using the approved typography roles. The reader PDF is
rendered from the same browser rulebook and shared design tokens used online,
then imposed into the preferred duplex booklet PDF.
"""

from __future__ import annotations

import argparse
import shutil
import socket
import subprocess
import sys
import tempfile
import time
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
RELEASE = ROOT / "releases" / "v0.6.1"
RULEBOOK_MD = RELEASE / "Gauntlet_v0.6.1_Rulebook.md"
REFERENCE_MD = RELEASE / "Gauntlet_v0.6.1_Reference_Guide.md"
RULEBOOK_DOCX = RELEASE / "Gauntlet_v0.6.1_Rulebook.docx"
RULEBOOK_PDF = RELEASE / "Gauntlet_v0.6.1_Rulebook.pdf"
RULEBOOK_BOOKLET_PDF = RELEASE / "Gauntlet_v0.6.1_Rulebook_Booklet.pdf"
REFERENCE_DOCX = RELEASE / "Gauntlet_v0.6.1_Reference_Guide.docx"
REFERENCE_PDF = RELEASE / "Gauntlet_v0.6.1_Reference_Guide.pdf"

ACCENT = RGBColor(95, 20, 24)
INK = RGBColor(24, 22, 20)
MUTED = RGBColor(102, 95, 85)

FONT_TITLE = "P22 1722 Pro"
FONT_HEADING = "Georgia"
FONT_READING = "Adobe Caslon Pro"
FONT_INTERFACE = "Inter"


def require(command: str) -> str:
    path = shutil.which(command)
    if not path:
        raise RuntimeError(f"Required command not found: {command}")
    return path


def set_run_font(run, name: str, size: float, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def ensure_style(document: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, style_type)


def build_reference_doc(
    path: Path,
    short_title: str,
    *,
    width: float,
    height: float,
    body_size: float,
    margin: float,
) -> None:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(width)
    section.page_height = Inches(height)
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)

    normal = document.styles["Normal"]
    normal.font.name = FONT_READING
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_READING)
    normal.font.size = Pt(body_size)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4.8)
    normal.paragraph_format.line_spacing = 1.04

    title = document.styles["Title"]
    title.font.name = FONT_TITLE
    title._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_TITLE)
    title.font.size = Pt(28 if width >= 8 else 24)
    title.font.bold = False
    title.font.color.rgb = ACCENT
    title.paragraph_format.space_after = Pt(9)

    subtitle = ensure_style(document, "Subtitle")
    subtitle.font.name = FONT_INTERFACE
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_INTERFACE)
    subtitle.font.size = Pt(10.5)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(9)

    for index, (size, before, after) in enumerate(
        [(18 if width >= 8 else 16, 14, 6), (14 if width >= 8 else 12.5, 10, 4), (11, 8, 3)],
        start=1,
    ):
        style = document.styles[f"Heading {index}"]
        style.font.name = FONT_HEADING if index < 3 else FONT_READING
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = index != 1
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = document.styles[style_name]
        style.font.name = FONT_READING
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_READING)
        style.font.size = Pt(body_size)
        style.paragraph_format.space_after = Pt(2.2)

    quote = document.styles["Quote"]
    quote.font.name = FONT_READING
    quote._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_READING)
    quote.font.size = Pt(body_size)
    quote.font.italic = True
    quote.font.color.rgb = INK
    quote.paragraph_format.left_indent = Inches(0.2)
    quote.paragraph_format.right_indent = Inches(0.12)
    quote.paragraph_format.space_before = Pt(4)
    quote.paragraph_format.space_after = Pt(5)

    table = ensure_style(document, "Table", WD_STYLE_TYPE.TABLE)
    table.font.name = FONT_INTERFACE
    table._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_INTERFACE)
    table.font.size = Pt(max(7.6, body_size - 1.3))

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(f"GAUNTLET  |  {short_title.upper()}  |  v0.6.1")
    set_run_font(run, FONT_INTERFACE, 7.2, bold=True)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Gauntlet v0.6.1  |  Page ")
    set_run_font(run, FONT_INTERFACE, 7.2)
    run.font.color.rgb = MUTED
    add_page_number(footer)

    document.add_paragraph("Reference template - generated; do not edit.")
    document.save(path)


def run(command: list[str], cwd: Path | None = None) -> None:
    print("+", " ".join(command), flush=True)
    subprocess.run(command, cwd=cwd, check=True)


def convert_markdown(
    markdown: Path,
    output_docx: Path,
    output_pdf: Path | None,
    title: str,
    reference_doc: Path,
    include_toc: bool,
) -> None:
    pandoc = require("pandoc")
    command = [
        pandoc,
        str(markdown),
        "--from=gfm",
        "--to=docx",
        f"--reference-doc={reference_doc}",
        f"--metadata=title:{title}",
        "--metadata=subtitle:Version 0.6.1 - First Playtest Revision",
        "--standalone",
        "--wrap=none",
        "-o",
        str(output_docx),
    ]
    if include_toc:
        command.extend(["--toc", "--toc-depth=2"])
    run(command, cwd=ROOT)

    if include_toc:
        run(
            [sys.executable, str(ROOT / "scripts/finalize_v061_docx.py"), str(output_docx)],
            cwd=ROOT,
        )

    if output_pdf is None:
        return

    libreoffice = require("libreoffice")
    with tempfile.TemporaryDirectory(prefix="gauntlet-lo-") as profile:
        run(
            [
                libreoffice,
                "--headless",
                f"-env:UserInstallation=file://{profile}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_pdf.parent),
                str(output_docx),
            ],
            cwd=ROOT,
        )
    generated = output_pdf.parent / f"{output_docx.stem}.pdf"
    if generated != output_pdf:
        generated.replace(output_pdf)


def wait_for_server(port: int, process: subprocess.Popen, timeout: float = 12.0) -> None:
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        if process.poll() is not None:
            raise RuntimeError("Local rulebook server stopped unexpectedly")
        with socket.socket() as client:
            client.settimeout(0.25)
            try:
                client.connect(("127.0.0.1", port))
                return
            except OSError:
                time.sleep(0.1)
    raise RuntimeError("Timed out waiting for local rulebook server")


def render_rulebook_pdfs() -> None:
    require("node")
    port = 8765
    server = subprocess.Popen(
        [sys.executable, "-m", "http.server", str(port), "--bind", "127.0.0.1"],
        cwd=ROOT,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    try:
        wait_for_server(port, server)
        run(
            [
                "node",
                str(ROOT / "scripts/render_v061_rulebook_pdf.mjs"),
                f"http://127.0.0.1:{port}/rulebook/",
                str(RULEBOOK_PDF),
            ],
            cwd=ROOT,
        )
    finally:
        server.terminate()
        try:
            server.wait(timeout=4)
        except subprocess.TimeoutExpired:
            server.kill()

    run(
        [
            sys.executable,
            str(ROOT / "scripts/impose_v061_rulebook_booklet.py"),
            str(RULEBOOK_PDF),
            str(RULEBOOK_BOOKLET_PDF),
        ],
        cwd=ROOT,
    )


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--check-tools", action="store_true")
    args = parser.parse_args()

    for command in ("pandoc", "libreoffice", "node"):
        require(command)
    if args.check_tools:
        print("Document build tools are available.")
        return 0

    if not RULEBOOK_MD.is_file() or not REFERENCE_MD.is_file():
        raise RuntimeError("Missing v0.6.1 Markdown sources")

    run([sys.executable, str(ROOT / "scripts/sync_v061_rulebook.py")], cwd=ROOT)

    RELEASE.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="gauntlet-docx-") as temp:
        temp_dir = Path(temp)
        rulebook_template = temp_dir / "rulebook-reference.docx"
        reference_template = temp_dir / "guide-reference.docx"
        build_reference_doc(
            rulebook_template,
            "Official Rulebook",
            width=5.5,
            height=8.5,
            body_size=9.4,
            margin=0.46,
        )
        build_reference_doc(
            reference_template,
            "Reference Guide",
            width=8.5,
            height=11,
            body_size=10.2,
            margin=0.7,
        )

        convert_markdown(
            RULEBOOK_MD,
            RULEBOOK_DOCX,
            None,
            "GAUNTLET - Official Rulebook",
            rulebook_template,
            include_toc=True,
        )
        convert_markdown(
            REFERENCE_MD,
            REFERENCE_DOCX,
            REFERENCE_PDF,
            "GAUNTLET - Reference Guide",
            reference_template,
            include_toc=False,
        )

    render_rulebook_pdfs()

    for output in [
        RULEBOOK_DOCX,
        RULEBOOK_PDF,
        RULEBOOK_BOOKLET_PDF,
        REFERENCE_DOCX,
        REFERENCE_PDF,
    ]:
        if not output.is_file() or output.stat().st_size < 10_000:
            raise RuntimeError(f"Output missing or unexpectedly small: {output}")
        print(f"Wrote {output.relative_to(ROOT)} ({output.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
