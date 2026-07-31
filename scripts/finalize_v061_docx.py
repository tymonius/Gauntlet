#!/usr/bin/env python3
"""Finalize the generated v0.6.1 Rulebook DOCX before PDF conversion.

Pandoc writes its DOCX table of contents as a dirty Word field. Headless
LibreOffice does not update that field before PDF conversion, producing an empty
contents section. This postprocessor replaces the field with a static linked
list built from the Rulebook's introductory and numbered section headings,
removes the redundant Markdown title block already represented by the DOCX
cover metadata, and places the approved hero sketch on a dedicated front cover.
"""

from __future__ import annotations

import argparse
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.image.image import Image
from docx.shared import Inches, Pt
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}
DOCUMENT_XML = "word/document.xml"
ROOT = Path(__file__).resolve().parents[1]
HERO_IMAGE = ROOT / "images/sketches/hero sketch.png"
HERO_ALT = "Gauntlet hero sketch"

INTRODUCTORY_HEADINGS = {"Welcome to Gauntlet", "Rules Conventions"}
NUMBERED_SECTION = re.compile(r"^\d+\.\s+")


def qn(local: str) -> str:
    return f"{{{W_NS}}}{local}"


def paragraph_text(element: etree._Element) -> str:
    return "".join(element.xpath(".//w:t/text()", namespaces=NS)).strip()


def paragraph_style(element: etree._Element) -> str:
    values = element.xpath("./w:pPr/w:pStyle/@w:val", namespaces=NS)
    return values[0] if values else ""


def nearest_bookmark(children: list[etree._Element], index: int) -> str | None:
    for previous in reversed(children[:index]):
        if previous.tag == qn("p"):
            break
        if previous.tag == qn("bookmarkStart"):
            name = previous.get(qn("name"), "")
            if name and not name.startswith("_"):
                return name
    return None


def make_paragraph(text: str, style: str, anchor: str | None = None) -> etree._Element:
    paragraph = etree.Element(qn("p"))
    properties = etree.SubElement(paragraph, qn("pPr"))
    style_node = etree.SubElement(properties, qn("pStyle"))
    style_node.set(qn("val"), style)

    parent = paragraph
    if anchor:
        hyperlink = etree.SubElement(paragraph, qn("hyperlink"))
        hyperlink.set(qn("anchor"), anchor)
        hyperlink.set(qn("history"), "1")
        parent = hyperlink

    run = etree.SubElement(parent, qn("r"))
    text_node = etree.SubElement(run, qn("t"))
    text_node.text = text
    return paragraph


def page_break_paragraph() -> etree._Element:
    paragraph = etree.Element(qn("p"))
    run = etree.SubElement(paragraph, qn("r"))
    break_node = etree.SubElement(run, qn("br"))
    break_node.set(qn("type"), "page")
    return paragraph


def is_contents_heading(style: str, text: str) -> bool:
    if style not in {"Heading1", "Heading2"}:
        return False
    return text in INTRODUCTORY_HEADINGS or bool(NUMBERED_SECTION.match(text))


def replace_toc(body: etree._Element) -> None:
    children = list(body)
    toc = next(
        (
            child
            for child in children
            if child.tag == qn("sdt")
            and child.xpath(
                ".//w:docPartGallery[@w:val='Table of Contents']",
                namespaces=NS,
            )
        ),
        None,
    )
    if toc is None:
        raise RuntimeError("Generated DOCX does not contain a Pandoc table-of-contents field")

    headings: list[tuple[str, str | None]] = []
    for index, child in enumerate(children):
        if child.tag != qn("p"):
            continue
        text = paragraph_text(child)
        style = paragraph_style(child)
        if not text or not is_contents_heading(style, text):
            continue
        headings.append((text, nearest_bookmark(children, index)))

    if len(headings) < 17:
        raise RuntimeError(
            f"Only {len(headings)} Rulebook section headings were found for the contents page"
        )

    content = toc.find("w:sdtContent", namespaces=NS)
    if content is None:
        raise RuntimeError("Generated DOCX table of contents has no content container")
    for child in list(content):
        content.remove(child)

    content.append(make_paragraph("Table of Contents", "TOCHeading"))
    for text, anchor in headings:
        content.append(make_paragraph(text, "TOC1", anchor))
    content.append(page_break_paragraph())


def remove_duplicate_markdown_title(body: etree._Element) -> None:
    children = list(body)
    start_index = next(
        (
            index
            for index, child in enumerate(children)
            if child.tag == qn("bookmarkStart") and child.get(qn("name")) == "gauntlet"
        ),
        None,
    )
    if start_index is None:
        return

    end_index = next(
        (
            index
            for index in range(start_index + 1, len(children))
            if children[index].tag == qn("bookmarkStart")
            and children[index].get(qn("name")) == "welcome-to-gauntlet"
        ),
        None,
    )
    if end_index is None:
        raise RuntimeError(
            "Could not locate the Welcome to Gauntlet bookmark after the duplicate title block"
        )

    for child in children[start_index:end_index]:
        body.remove(child)


def insert_cover_image(path: Path) -> None:
    if not HERO_IMAGE.is_file() or HERO_IMAGE.stat().st_size == 0:
        raise RuntimeError(f"Rulebook hero sketch is missing: {HERO_IMAGE.relative_to(ROOT)}")

    document = Document(path)
    if HERO_ALT in document._element.xml:
        return

    subtitle = next(
        (
            paragraph
            for paragraph in document.paragraphs
            if paragraph.style.name == "Subtitle"
            or paragraph.text.startswith("Version 0.6.1")
        ),
        None,
    )
    if subtitle is None:
        raise RuntimeError("Could not locate the DOCX cover subtitle")

    source = Image.from_file(str(HERO_IMAGE))
    max_width = Inches(4.25)
    max_height = Inches(4.75)
    scale = min(float(max_width) / float(source.width), float(max_height) / float(source.height))
    width = int(float(source.width) * scale)
    height = int(float(source.height) * scale)

    cover = document.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(14)
    cover.paragraph_format.space_after = Pt(0)
    cover.paragraph_format.keep_together = True
    picture_run = cover.add_run()
    picture_run.add_picture(str(HERO_IMAGE), width=width, height=height)
    for properties in picture_run._r.xpath(".//pic:cNvPr"):
        properties.set("descr", HERO_ALT)
        properties.set("name", HERO_ALT)
    cover.add_run().add_break(WD_BREAK.PAGE)
    subtitle._p.addnext(cover._p)

    section = document.sections[0]
    section.different_first_page_header_footer = True
    for paragraph in section.first_page_header.paragraphs:
        paragraph.clear()
    for paragraph in section.first_page_footer.paragraphs:
        paragraph.clear()

    document.save(path)


def finalize(path: Path) -> None:
    if not path.is_file():
        raise FileNotFoundError(path)

    with zipfile.ZipFile(path, "r") as source:
        document_xml = source.read(DOCUMENT_XML)
        root = etree.fromstring(document_xml)
        body = root.find("w:body", namespaces=NS)
        if body is None:
            raise RuntimeError("Generated DOCX has no document body")

        replace_toc(body)
        remove_duplicate_markdown_title(body)
        updated_xml = etree.tostring(
            root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )

        with tempfile.NamedTemporaryFile(
            prefix=f"{path.stem}-",
            suffix=".docx",
            dir=path.parent,
            delete=False,
        ) as temporary:
            temporary_path = Path(temporary.name)

        try:
            with zipfile.ZipFile(temporary_path, "w") as target:
                for item in source.infolist():
                    payload = updated_xml if item.filename == DOCUMENT_XML else source.read(item.filename)
                    target.writestr(item, payload)
            temporary_path.replace(path)
        finally:
            temporary_path.unlink(missing_ok=True)

    insert_cover_image(path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path, help="Generated Rulebook DOCX to finalize in place")
    args = parser.parse_args()
    finalize(args.docx.resolve())
    print(f"Finalized illustrated Rulebook cover and visible contents in {args.docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
