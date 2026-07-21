#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path('build/v0.6.0/reference.docx')
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.68)
section.bottom_margin = Inches(0.68)
section.left_margin = Inches(0.72)
section.right_margin = Inches(0.72)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Noto Serif'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.08

for name, size, before, after in [
    ('Title', 34, 0, 12),
    ('Subtitle', 16, 0, 10),
    ('Heading 1', 22, 16, 8),
    ('Heading 2', 15, 12, 5),
    ('Heading 3', 12, 9, 3),
    ('Heading 4', 10.5, 6, 2),
]:
    style = styles[name]
    style.font.name = 'Noto Sans'
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(23, 34, 53)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

styles['Title'].font.bold = True
styles['Title'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
styles['Subtitle'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

if 'Compact Table' not in styles:
    table_style = styles.add_style('Compact Table', WD_STYLE_TYPE.TABLE)
else:
    table_style = styles['Compact Table']
table_style.font.name = 'Noto Sans'
table_style.font.size = Pt(9)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header.add_run('GAUNTLET  •  OFFICIAL RULEBOOK  •  v0.6.0')
run.font.name = 'Noto Sans'
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(100, 116, 139)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— ')
run.font.name = 'Noto Sans'
run.font.size = Pt(8)

field = OxmlElement('w:fldSimple')
field.set(qn('w:instr'), 'PAGE')
footer._p.append(field)
footer.add_run(' —')

# Seed paragraphs ensure Pandoc retains the configured title styles and page geometry.
doc.add_paragraph('GAUNTLET', style='Title')
doc.add_paragraph('Official Rulebook', style='Subtitle')
doc.add_section(WD_SECTION_START.NEW_PAGE)

doc.save(OUTPUT)
print(OUTPUT)
