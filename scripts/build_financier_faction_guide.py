from pathlib import Path

import mistune
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
GUIDE_DIR = ROOT / "releases/v0.6/faction-guides/financier"
MD_PATH = GUIDE_DIR / "Gauntlet_v0.6_Financier_Faction_Guide.md"
OUT = GUIDE_DIR / "Gauntlet_v0.6_Financier_Faction_Guide.docx"
SKETCHES = ROOT / "images/sketches"


def prepare_leader_sketches():
    SKETCHES.mkdir(parents=True, exist_ok=True)
    crops = {
        "banker": (90, 70, 500, 650),
        "executive": (300, 25, 790, 750),
    }
    for name, box in crops.items():
        source = ROOT / "images" / f"{name}.png"
        destination = SKETCHES / f"{name}.png"
        with Image.open(source) as image:
            image.crop(box).save(destination)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_bottom_border(paragraph, color="7A5D1B", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_page_field(run):
    start = OxmlElement("w:fldChar")
    start.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(start)
    run._r.append(instruction)
    run._r.append(end)


def initialize_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.23)
    section.footer_distance = Inches(0.23)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(8.5)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0

    title = styles["Title"]
    title.font.name = "Georgia"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(91, 69, 19)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True

    heading_specs = [
        ("Heading 1", 16.5, "5B4513", 10, 5),
        ("Heading 2", 12.5, "7A5D1B", 8, 3),
        ("Heading 3", 10.0, "5B4513", 6, 2),
        ("Heading 4", 9.0, "5B4513", 4, 2),
    ]
    for name, size, color, before, after in heading_specs:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    quote = styles["Quote"]
    quote.font.name = "Aptos"
    quote.font.size = Pt(8.5)
    quote.font.italic = False
    quote.font.color.rgb = RGBColor(45, 45, 45)
    quote.paragraph_format.left_indent = Inches(0.18)
    quote.paragraph_format.right_indent = Inches(0.08)
    quote.paragraph_format.space_before = Pt(3)
    quote.paragraph_format.space_after = Pt(4)
    p_pr = quote.element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), "B08A32")
    borders.append(left)
    p_pr.append(borders)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(8.5)
        style.paragraph_format.space_after = Pt(1.5)

    header = section.header.paragraphs[0]
    header.text = "GAUNTLET v0.6  |  FINANCIER FACTION GUIDE"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.runs[0]
    header_run.font.name = "Aptos"
    header_run.font.size = Pt(7)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(95, 78, 37)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("FINANCIERS  •  ")
    footer_run.font.name = "Aptos"
    footer_run.font.size = Pt(7)
    footer_run.font.color.rgb = RGBColor(95, 78, 37)
    add_page_field(footer_run)
    return document


def text_from_nodes(nodes):
    output = []
    for node in nodes or []:
        node_type = node.get("type")
        if node_type in ("text", "codespan", "inline_html"):
            output.append(node.get("raw", ""))
        elif node_type in ("strong", "emphasis", "link"):
            output.append(text_from_nodes(node.get("children", [])))
        elif node_type in ("softbreak", "linebreak"):
            output.append("\n")
        elif node_type == "image":
            output.append(text_from_nodes(node.get("children", [])))
    return "".join(output)


def add_inline(paragraph, nodes, bold=False, italic=False):
    for node in nodes or []:
        node_type = node.get("type")
        if node_type == "text":
            run = paragraph.add_run(node.get("raw", ""))
            run.bold = bold
            run.italic = italic
        elif node_type == "strong":
            add_inline(paragraph, node.get("children", []), True, italic)
        elif node_type == "emphasis":
            add_inline(paragraph, node.get("children", []), bold, True)
        elif node_type == "codespan":
            run = paragraph.add_run(node.get("raw", ""))
            run.bold = bold
            run.italic = italic
            run.font.name = "Consolas"
            run.font.size = Pt(7.7)
        elif node_type in ("softbreak", "linebreak"):
            paragraph.add_run().add_break()
        elif node_type == "link":
            add_inline(paragraph, node.get("children", []), bold, italic)
        elif node_type == "image":
            alt = text_from_nodes(node.get("children", [])).lower()
            image = SKETCHES / ("banker.png" if "banker" in alt else "executive.png")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.add_run().add_picture(str(image), width=Inches(2.05))
        elif node_type == "inline_html":
            raw = node.get("raw", "")
            if raw.strip() not in ("<br>", "<br/>", "<br />"):
                run = paragraph.add_run(raw)
                run.bold = bold
                run.italic = italic


def render_blocks(document, blocks, quote=False, list_level=0):
    for node in blocks:
        node_type = node.get("type")
        if node_type == "blank_line":
            continue
        if node_type == "heading":
            level = node.get("attrs", {}).get("level", 1)
            style = "Title" if level == 1 else f"Heading {min(level - 1, 4)}"
            paragraph = document.add_paragraph(style=style)
            add_inline(paragraph, node.get("children", []))
        elif node_type in ("paragraph", "block_text"):
            paragraph = document.add_paragraph(style="Quote" if quote else None)
            add_inline(paragraph, node.get("children", []))
        elif node_type == "block_quote":
            render_blocks(document, node.get("children", []), quote=True, list_level=list_level)
        elif node_type == "thematic_break":
            add_bottom_border(document.add_paragraph())
        elif node_type == "list":
            ordered = node.get("attrs", {}).get("ordered", False)
            style = "List Number" if ordered else "List Bullet"
            for item in node.get("children", []):
                first = True
                for child in item.get("children", []):
                    if child.get("type") in ("block_text", "paragraph"):
                        paragraph = document.add_paragraph(style=style if first else None)
                        if list_level:
                            paragraph.paragraph_format.left_indent = Inches(0.25 + 0.18 * list_level)
                        add_inline(paragraph, child.get("children", []))
                        first = False
                    elif child.get("type") == "list":
                        render_blocks(document, [child], quote=quote, list_level=list_level + 1)
        elif node_type == "table":
            head = []
            rows = []
            for part in node.get("children", []):
                if part.get("type") == "table_head":
                    head = [cell.get("children", []) for cell in part.get("children", [])]
                elif part.get("type") == "table_body":
                    for row in part.get("children", []):
                        rows.append([cell.get("children", []) for cell in row.get("children", [])])
            column_count = max([len(head)] + [len(row) for row in rows] + [1])
            table = document.add_table(rows=(1 if head else 0) + len(rows), cols=column_count)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            row_offset = 0
            if head:
                header_row = table.rows[0]
                set_repeat_table_header(header_row)
                for index, nodes in enumerate(head):
                    cell = header_row.cells[index]
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_shading(cell, "E9DFC4")
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_inline(paragraph, nodes, bold=True)
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(91, 69, 19)
                row_offset = 1
            for row_index, row_nodes in enumerate(rows):
                row = table.rows[row_offset + row_index]
                for column_index, nodes in enumerate(row_nodes):
                    cell = row.cells[column_index]
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    add_inline(paragraph, nodes)
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
            if column_count == 2:
                for row in table.rows:
                    row.cells[0].width = Inches(1.65)
                    row.cells[1].width = Inches(5.55)
            document.add_paragraph().paragraph_format.space_after = Pt(0)
        elif node_type == "block_code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(node.get("raw", ""))
            run.font.name = "Consolas"
            run.font.size = Pt(7.5)
        elif node.get("children"):
            render_blocks(document, node.get("children", []), quote=quote, list_level=list_level)


def main():
    prepare_leader_sketches()
    markdown = mistune.create_markdown(renderer="ast", plugins=["table"])
    syntax_tree = markdown(MD_PATH.read_text(encoding="utf-8"))
    document = initialize_document()
    render_blocks(document, syntax_tree)
    for table in document.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            tr_pr.append(OxmlElement("w:cantSplit"))
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
